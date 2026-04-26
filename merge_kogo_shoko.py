#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
甲号証ファイル結合スクリプト
============================

複数の「甲第xxx号証.docx」ファイルを番号順に結合し、結合甲号証ファイル
（綴じ込みファイル）を生成します。

【主な機能】
1. ファイル名・本文中の【甲第xxx号証】マーカーから番号を抽出
   - 半角／全角／スペース／「第」の有無 などの表記ゆれを吸収
   - 枝番（例：その１、その2）も認識
2. 番号を「全角数字3桁」に正規化（例：1 → ００１ 、12 → ０１２、123 → １２３）
3. 番号順（枝番がある場合は枝番順）でソート
4. 各ファイルの本文先頭マーカーを正規化形式へ書き換え
5. python-docx + docxcompose で結合（書式・画像・表をできる限り保持）
6. ページ区切りで各甲号証の境界を明示
7. 出力ファイル名も全角3桁で統一

【依存ライブラリ】
    pip install python-docx docxcompose

【使い方】
    python merge_kogo_shoko.py 入力ファイル1.docx 入力ファイル2.docx ...
    python merge_kogo_shoko.py --indir ./入力フォルダ --outdir ./出力フォルダ

"""

from __future__ import annotations

import argparse
import re
import shutil
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Tuple, List

from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer

# ----------------------------------------------------------------------
# 番号正規化ユーティリティ
# ----------------------------------------------------------------------

# 全角数字 → 半角数字 変換テーブル
ZEN2HAN_DIGITS = str.maketrans("０１２３４５６７８９", "0123456789")
# 半角数字 → 全角数字 変換テーブル
HAN2ZEN_DIGITS = str.maketrans("0123456789", "０１２３４５６７８９")

def to_han_digits(s: str) -> str:
    """全角数字を半角数字に変換"""
    return s.translate(ZEN2HAN_DIGITS)

def to_zen_digits(s: str) -> str:
    """半角数字を全角数字に変換"""
    return s.translate(HAN2ZEN_DIGITS)

def normalize_main_number(num_str: str, width: int = 3) -> str:
    """
    主番号を全角3桁にゼロ埋めして返す。
    例: "1" → "００１", "12" → "０１２", "123" → "１２３", "1234" → "１２３４"
    """
    han = to_han_digits(num_str)
    n = int(han)
    # 3桁ゼロ埋め（4桁以上の場合はそのまま桁数維持）
    padded = str(n).zfill(width)
    return to_zen_digits(padded)

def normalize_branch_number(num_str: str) -> str:
    """枝番を全角に変換（ゼロ埋めなし、人間可読性優先）"""
    han = to_han_digits(num_str)
    n = int(han)
    return to_zen_digits(str(n))

# ----------------------------------------------------------------------
# 番号抽出（正規表現）
# ----------------------------------------------------------------------

# マーカー全体を捉える：【甲第 ＸＸ 号証(その Ｙ)】
# - "甲" の後の "第" は省略可
# - 数字は半角／全角混在可、前後にスペース可
# - 枝番 "その" の後の数字も同様
MARKER_PATTERN = re.compile(
    r"【?\s*甲\s*第?\s*"                          # 「【甲第」
    r"(?P<main>[0-9０-９]+)"                       # 主番号
    r"\s*号\s*証"                                  # 「号証」
    r"(?:\s*その\s*(?P<branch>[0-9０-９]+))?"      # （枝番）
    r"\s*】?"
)

# ファイル名用の緩めのパターン（拡張子前まで）
FILENAME_PATTERN = re.compile(
    r"甲\s*第?\s*"
    r"(?P<main>[0-9０-９]+)"
    r"\s*号\s*証"
    r"(?:\s*その\s*(?P<branch>[0-9０-９]+))?"
)

@dataclass
class KogoNumber:
    """甲号証番号（正規化済み）"""
    main: int                       # 主番号（整数）
    branch: Optional[int] = None    # 枝番（無ければ None）

    @property
    def sort_key(self) -> Tuple[int, int]:
        """ソート用キー：主番号 → 枝番"""
        return (self.main, self.branch if self.branch is not None else 0)

    @property
    def normalized_marker(self) -> str:
        """正規化された本文マーカー文字列"""
        main_str = normalize_main_number(str(self.main))
        if self.branch is not None:
            return f"【甲第{main_str}号証その{normalize_branch_number(str(self.branch))}】"
        return f"【甲第{main_str}号証】"

    @property
    def normalized_filename_stem(self) -> str:
        """正規化されたファイル名（拡張子なし）"""
        main_str = normalize_main_number(str(self.main))
        if self.branch is not None:
            return f"甲第{main_str}号証その{normalize_branch_number(str(self.branch))}"
        return f"甲第{main_str}号証"

def extract_number_from_text(text: str) -> Optional[KogoNumber]:
    """テキストから甲号証番号を抽出"""
    m = MARKER_PATTERN.search(text)
    if not m:
        return None
    main = int(to_han_digits(m.group("main")))
    branch_raw = m.group("branch")
    branch = int(to_han_digits(branch_raw)) if branch_raw else None
    return KogoNumber(main=main, branch=branch)

def extract_number_from_filename(path: Path) -> Optional[KogoNumber]:
    """ファイル名から甲号証番号を抽出"""
    m = FILENAME_PATTERN.search(path.stem)
    if not m:
        return None
    main = int(to_han_digits(m.group("main")))
    branch_raw = m.group("branch")
    branch = int(to_han_digits(branch_raw)) if branch_raw else None
    return KogoNumber(main=main, branch=branch)

def detect_number(path: Path) -> KogoNumber:
    """
    ファイル名 と 文書冒頭マーカー の両方から番号を抽出。
    本文中のマーカーを優先（より正確）。一致しない場合は警告を出して本文側を採用。
    """
    from_doc = None
    try:
        doc = Document(str(path))
        # 先頭から最大10段落を走査
        for para in doc.paragraphs[:10]:
            from_doc = extract_number_from_text(para.text)
            if from_doc:
                break
    except Exception as e:
        print(f"  [警告] 本文読込失敗: {e}", file=sys.stderr)

    from_name = extract_number_from_filename(path)

    if from_doc and from_name and from_doc.sort_key != from_name.sort_key:
        print(
            f"  [警告] {path.name}: ファイル名({from_name.sort_key}) と "
            f"本文マーカー({from_doc.sort_key}) が不一致 → 本文を優先",
            file=sys.stderr,
        )

    if from_doc:
        return from_doc
    if from_name:
        return from_name
    raise ValueError(f"{path.name}: 甲号証番号を抽出できませんでした")

# ----------------------------------------------------------------------
# 本文マーカー書き換え
# ----------------------------------------------------------------------

def rewrite_marker_in_document(doc_path: Path, kogo: KogoNumber, out_path: Path) -> None:
    """
    文書冒頭の【甲第xxx号証】マーカーを正規化形式に書き換えて保存する。

    本文中のマーカーは複数 run に分かれている可能性があるため、
    対象段落を特定したら、最初の run にまとめてテキストを入れ、
    それ以外の run のテキストは空にする（書式は最初の run のものを保持）。
    """
    doc = Document(str(doc_path))
    rewritten = False

    for para in doc.paragraphs[:10]:  # 冒頭付近のみ走査
        if MARKER_PATTERN.search(para.text):
            # 段落全体のテキストを正規化マーカーで置き換える
            new_text = MARKER_PATTERN.sub(kogo.normalized_marker, para.text, count=1)

            if not para.runs:
                # run が無い段落は新規追加
                para.add_run(new_text)
            else:
                # 最初の run にまとめて格納（書式維持）
                para.runs[0].text = new_text
                # 残りの run のテキストをクリア
                for r in para.runs[1:]:
                    r.text = ""
            rewritten = True
            break

    if not rewritten:
        # 段落本文に見つからない場合でも、ファイルはコピーして処理続行
        print(f"  [情報] {doc_path.name}: 本文マーカーは見つからず（書き換え省略）", file=sys.stderr)

    doc.save(str(out_path))

# ----------------------------------------------------------------------
# 結合処理（docxcompose）
# ----------------------------------------------------------------------

def insert_pagebreak_at_end(doc_path: Path) -> None:
    """文書末尾に改ページを追加（綴じ込み時の境界を明確化）"""
    doc = Document(str(doc_path))
    # 末尾段落に改ページ run を追加
    last_para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    run = last_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    doc.save(str(doc_path))

def merge_documents(prepared_files: List[Path], output_path: Path) -> None:
    """
    docxcompose で複数 docx を結合する。
    最初のファイルをベースに、残りを順次追加する方式。
    書式・画像・スタイル・番号定義などを高い忠実度で保持。
    """
    if not prepared_files:
        raise ValueError("結合対象ファイルがありません")

    base = Document(str(prepared_files[0]))
    composer = Composer(base)
    for f in prepared_files[1:]:
        composer.append(Document(str(f)))
    composer.save(str(output_path))

# ----------------------------------------------------------------------
# メイン
# ----------------------------------------------------------------------

def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description="甲号証 docx ファイルを番号順に結合します。",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("inputs", nargs="*", type=Path, help="入力 docx ファイル")
    parser.add_argument("--indir", type=Path, help="入力ディレクトリ（中の *.docx を一括処理）")
    parser.add_argument(
        "-o", "--output", type=Path,
        default=Path("結合甲号証.docx"),
        help="出力ファイル名（デフォルト: 結合甲号証.docx）",
    )
    parser.add_argument(
        "--rename", action="store_true",
        help="入力ファイルのコピーを正規化済みファイル名で別フォルダに保存する",
    )
    parser.add_argument(
        "--renamed-dir", type=Path, default=Path("正規化済み"),
        help="--rename 指定時の保存先ディレクトリ",
    )
    parser.add_argument(
        "--no-pagebreak", action="store_true",
        help="各甲号証の末尾に改ページを挿入しない",
    )
    args = parser.parse_args(argv)

    # 入力ファイル収集
    inputs: List[Path] = list(args.inputs)
    if args.indir:
        inputs.extend(sorted(args.indir.glob("*.docx")))
    # 一時ファイルを除外
    inputs = [p for p in inputs if not p.name.startswith("~$")]

    if not inputs:
        parser.error("入力ファイルが指定されていません")

    # 番号を抽出してソート
    print("[1/3] 甲号証番号を抽出中...")
    pairs: List[Tuple[KogoNumber, Path]] = []
    for p in inputs:
        try:
            kogo = detect_number(p)
            print(f"  {p.name}  →  {kogo.normalized_marker}")
            pairs.append((kogo, p))
        except ValueError as e:
            print(f"  [スキップ] {e}", file=sys.stderr)

    pairs.sort(key=lambda x: x[0].sort_key)

    # 重複チェック
    seen = set()
    for kogo, p in pairs:
        if kogo.sort_key in seen:
            print(f"  [警告] 番号重複: {kogo.normalized_marker} ({p.name})", file=sys.stderr)
        seen.add(kogo.sort_key)

    # 各ファイルを正規化（マーカー書き換え）→一時ディレクトリへ
    print("\n[2/3] 本文マーカーを正規化中...")
    workdir = Path(tempfile.mkdtemp(prefix="kogo_merge_"))
    prepared: List[Path] = []
    try:
        for idx, (kogo, src) in enumerate(pairs):
            stem = kogo.normalized_filename_stem
            tmp_path = workdir / f"{idx:03d}_{stem}.docx"
            rewrite_marker_in_document(src, kogo, tmp_path)

            # 末尾改ページ（最後のファイル以外）
            if not args.no_pagebreak and idx < len(pairs) - 1:
                insert_pagebreak_at_end(tmp_path)
            prepared.append(tmp_path)
            print(f"  {src.name}  →  {stem}.docx")

            # ファイル名正規化コピー
            if args.rename:
                args.renamed_dir.mkdir(parents=True, exist_ok=True)
                shutil.copy2(src, args.renamed_dir / f"{stem}.docx")

        # 結合
        print(f"\n[3/3] 結合中 → {args.output}")
        args.output.parent.mkdir(parents=True, exist_ok=True)
        merge_documents(prepared, args.output)
        print("完了。")

    finally:
        shutil.rmtree(workdir, ignore_errors=True)

    return 0

if __name__ == "__main__":
    sys.exit(main())
