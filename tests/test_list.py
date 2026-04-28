# -*- coding: utf-8 -*-
"""
甲号証リスト生成（list_service）と /api/auto-list のテスト。
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.list_service import generate_list
from app.main import app
from app.merge_service import (
    LIST_FILENAME,
    MASTER_DIRNAME,
    OUTPUT_DIRNAME,
    OUTPUT_FILENAME,
    BACKUP_SUFFIX,
    ensure_folders,
    merge_kogo,
)

from tests.conftest import make_kogo_docx, make_list_docx


# ---------------------------------------------------------------------------
# ヘルパ
# ---------------------------------------------------------------------------

def _read_paragraphs(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    return [p.text for p in doc.paragraphs if p.text != ""]


@pytest.fixture
def client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
    monkeypatch.setenv("KOGO_KANRI_SETTINGS_PATH", str(tmp_path / "settings.json"))
    return TestClient(app)


# ---------------------------------------------------------------------------
# source="master" 正常系
# ---------------------------------------------------------------------------

def test_generate_list_master_writes_normalized_stems(root_folder: Path) -> None:
    ensure_folders(root_folder)
    master = root_folder / MASTER_DIRNAME
    make_kogo_docx(master / "甲第００１号証.docx", "【甲第１号証】")
    make_kogo_docx(master / "甲第００２号証.docx", "【甲第２号証】")
    make_kogo_docx(master / "甲第０１２号証その１.docx", "【甲第０１２号証その１】")

    # ensure_folders が空のリストを作るので、それは事前に削除しておく
    list_path = root_folder / LIST_FILENAME
    if list_path.exists():
        list_path.unlink()

    outcome = generate_list(root_folder, source="master")

    assert outcome.source == "master"
    assert outcome.output_path == list_path
    assert outcome.backup_created is False
    assert outcome.warnings == []
    # 並び順は (main, branch) 昇順、表記は正規化ファイル名形式（【】無し）
    assert outcome.numbers_written == [
        "甲第００１号証",
        "甲第００２号証",
        "甲第０１２号証その１",
    ]
    # 実ファイルにも書かれている
    paragraphs = _read_paragraphs(list_path)
    assert paragraphs == [
        "甲第００１号証",
        "甲第００２号証",
        "甲第０１２号証その１",
    ]
    # マーカー形式（【】）は使わない
    for line in paragraphs:
        assert "【" not in line and "】" not in line


def test_generate_list_master_sorts_main_branch(root_folder: Path) -> None:
    """ファイル名のアルファベット順ではなく (main, branch) 順で並ぶ。"""
    ensure_folders(root_folder)
    master = root_folder / MASTER_DIRNAME
    make_kogo_docx(master / "甲第０１０号証.docx", "【甲第１０号証】")
    make_kogo_docx(master / "甲第００２号証.docx", "【甲第２号証】")
    make_kogo_docx(master / "甲第００５号証.docx", "【甲第５号証】")
    (root_folder / LIST_FILENAME).unlink()

    outcome = generate_list(root_folder, source="master")

    assert outcome.numbers_written == [
        "甲第００２号証",
        "甲第００５号証",
        "甲第０１０号証",
    ]


# ---------------------------------------------------------------------------
# source="combined" 正常系
# ---------------------------------------------------------------------------

def test_generate_list_combined_writes_normalized_stems(root_folder: Path) -> None:
    ensure_folders(root_folder)
    master = root_folder / MASTER_DIRNAME
    make_kogo_docx(master / "甲第００１号証.docx", "【甲第１号証】")
    make_kogo_docx(master / "甲第００２号証.docx", "【甲第２号証】")
    make_kogo_docx(master / "甲第００３号証.docx", "【甲第３号証】")
    # 結合 docx を生成
    merge_kogo(root_folder)

    # 既存の甲号証リスト.docx は除けておく（バックアップ動作は別テストで検証）
    list_path = root_folder / LIST_FILENAME
    if list_path.exists():
        list_path.unlink()

    outcome = generate_list(root_folder, source="combined")

    assert outcome.source == "combined"
    assert outcome.output_path == list_path
    assert outcome.numbers_written == [
        "甲第００１号証",
        "甲第００２号証",
        "甲第００３号証",
    ]
    paragraphs = _read_paragraphs(list_path)
    assert paragraphs == [
        "甲第００１号証",
        "甲第００２号証",
        "甲第００３号証",
    ]


# ---------------------------------------------------------------------------
# 既存リストの .bak バックアップ
# ---------------------------------------------------------------------------

def test_generate_list_backs_up_existing(root_folder: Path) -> None:
    ensure_folders(root_folder)
    master = root_folder / MASTER_DIRNAME
    make_kogo_docx(master / "甲第００１号証.docx", "【甲第１号証】")

    # 既存の甲号証リスト.docx に固有の本文を入れておく
    list_path = root_folder / LIST_FILENAME
    make_list_docx(list_path, ["手動編集された行"])

    outcome = generate_list(root_folder, source="master")

    assert outcome.backup_created is True
    backup_path = list_path.with_suffix(list_path.suffix + BACKUP_SUFFIX)
    assert backup_path.exists()
    # バックアップには元の内容が残っている
    backup_paragraphs = _read_paragraphs(backup_path)
    assert "手動編集された行" in backup_paragraphs
    # 新しいリストには新内容が入る
    new_paragraphs = _read_paragraphs(list_path)
    assert new_paragraphs == ["甲第００１号証"]


# ---------------------------------------------------------------------------
# ルートフォルダ不在 → HTTP 400
# ---------------------------------------------------------------------------

def test_auto_list_endpoint_missing_root_returns_400(client: TestClient) -> None:
    r = client.post(
        "/api/auto-list",
        json={"root_folder": "/no/such/path/zzz", "source": "master"},
    )
    assert r.status_code == 400


# ---------------------------------------------------------------------------
# 結合甲号証不在 (source="combined") → HTTP 400
# ---------------------------------------------------------------------------

def test_auto_list_endpoint_combined_missing_returns_400(
    client: TestClient, tmp_path: Path
) -> None:
    root = tmp_path / "case_combined_missing"
    ensure_folders(root)
    # 結合 docx は生成しない
    r = client.post(
        "/api/auto-list",
        json={"root_folder": str(root), "source": "combined"},
    )
    assert r.status_code == 400
    assert "結合甲号証" in r.json()["detail"]


# ---------------------------------------------------------------------------
# 番号抽出できないファイルが警告に記録される
# ---------------------------------------------------------------------------

def test_generate_list_master_warns_on_unrecognized(root_folder: Path) -> None:
    ensure_folders(root_folder)
    master = root_folder / MASTER_DIRNAME
    make_kogo_docx(master / "甲第００１号証.docx", "【甲第１号証】")
    # 番号が抽出できない docx
    bad = master / "memo.docx"
    d = Document()
    d.add_paragraph("メモ")
    d.save(str(bad))
    (root_folder / LIST_FILENAME).unlink()

    outcome = generate_list(root_folder, source="master")

    # 抽出できたものは含まれ、できないものは除外
    assert outcome.numbers_written == ["甲第００１号証"]
    # 警告が記録されている
    assert outcome.warnings
    assert any("memo.docx" in w for w in outcome.warnings)


# ---------------------------------------------------------------------------
# 重複番号: 警告に記録 + 最初の1件のみ採用
# ---------------------------------------------------------------------------

def test_generate_list_master_dedupes_with_warning(tmp_path: Path) -> None:
    root = tmp_path / "case_dup"
    ensure_folders(root)
    master = root / MASTER_DIRNAME
    make_kogo_docx(master / "甲第００１号証.docx", "【甲第１号証】")
    make_kogo_docx(master / "甲第００１号証_別.docx", "【甲第１号証】")
    make_kogo_docx(master / "甲第００２号証.docx", "【甲第２号証】")
    (root / LIST_FILENAME).unlink()

    outcome = generate_list(root, source="master")

    assert outcome.numbers_written == ["甲第００１号証", "甲第００２号証"]
    assert any("重複" in w for w in outcome.warnings)


# ---------------------------------------------------------------------------
# /api/auto-list エンドポイント正常系
# ---------------------------------------------------------------------------

def test_auto_list_endpoint_master_happy_path(
    client: TestClient, tmp_path: Path
) -> None:
    root = tmp_path / "case_api_master"
    ensure_folders(root)
    master = root / MASTER_DIRNAME
    make_kogo_docx(master / "甲第００１号証.docx", "【甲第１号証】")
    make_kogo_docx(master / "甲第００２号証.docx", "【甲第２号証】")

    r = client.post(
        "/api/auto-list",
        json={"root_folder": str(root), "source": "master"},
    )
    assert r.status_code == 200
    d = r.json()
    assert d["source"] == "master"
    assert d["numbers_written"] == ["甲第００１号証", "甲第００２号証"]
    # ensure_folders で空リストが作られているので backup_created は True
    assert d["backup_created"] is True
    assert Path(d["output_path"]).exists()


def test_auto_list_endpoint_invalid_source_returns_400(
    client: TestClient, tmp_path: Path
) -> None:
    root = tmp_path / "case_invalid_source"
    ensure_folders(root)
    r = client.post(
        "/api/auto-list",
        json={"root_folder": str(root), "source": "bogus"},
    )
    assert r.status_code == 400
