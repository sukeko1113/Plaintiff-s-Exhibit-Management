# -*- coding: utf-8 -*-
"""
ai_client.py の単体テスト(SPEC §10.5、§10.8)。

実 API 呼び出しは行わず、`anthropic.Anthropic` クライアント相当の
モックオブジェクトを `extract_metadata(client=...)` に注入することで
SDK 全体をスタブ化する。
"""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import MagicMock
from typing import Any, List

import pytest
from docx import Document

import anthropic

from app import ai_client
from app.ai_client import (
    AiAuthError,
    AiClientError,
    AiParseError,
    AiRateLimitError,
    AiTimeoutError,
    ExtractedMetadata,
    extract_metadata,
    extract_text_from_docx,
    load_system_prompt,
    truncate_for_tokens,
)


# ---------------------------------------------------------------------------
# 共通ヘルパ
# ---------------------------------------------------------------------------

def _make_docx(path: Path, paragraphs: List[str], tables: List[List[List[str]]] | None = None) -> Path:
    """指定段落と表を持つ docx を生成。"""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if tables:
        for t in tables:
            rows = len(t)
            cols = len(t[0]) if t else 0
            tbl = doc.add_table(rows=rows, cols=cols)
            for r, row in enumerate(t):
                for c, val in enumerate(row):
                    tbl.rows[r].cells[c].text = val
    doc.save(str(path))
    return path


def _make_response(text: str) -> MagicMock:
    """anthropic.messages.create() の戻り値を模したモック。"""
    block = MagicMock()
    block.text = text
    resp = MagicMock()
    resp.content = [block]
    return resp


def _make_client(create_side_effect: Any) -> MagicMock:
    """messages.create に side_effect を設定したクライアントモック。"""
    client = MagicMock()
    client.messages = MagicMock()
    if isinstance(create_side_effect, list):
        client.messages.create.side_effect = create_side_effect
    else:
        client.messages.create.return_value = create_side_effect
    return client


def _raw_anthropic_error(cls):
    """anthropic SDK の Error クラスを __init__ をバイパスして生成。"""
    err = cls.__new__(cls)
    Exception.__init__(err, "test")
    return err


_VALID_JSON = json.dumps({
    "title": "業務委託契約書",
    "created_date": "令和5年4月1日",
    "author": "山田太郎",
    "purpose": "契約締結の事実を立証",
    "confidence": {
        "title": "high",
        "created_date": "high",
        "author": "medium",
        "purpose": "low",
    },
    "notes": "",
}, ensure_ascii=False)


def _long_body() -> str:
    """MIN_TEXT_LENGTH を超える適度な長さの本文。"""
    return "これは契約書の本文です。" * 30  # 約 360 文字


# ---------------------------------------------------------------------------
# テキスト抽出
# ---------------------------------------------------------------------------

def test_extract_text_collects_paragraphs_and_tables(tmp_path: Path) -> None:
    docx = _make_docx(
        tmp_path / "a.docx",
        ["第1段落", "第2段落"],
        tables=[[["セル1-1", "セル1-2"], ["セル2-1", "セル2-2"]]],
    )
    text = extract_text_from_docx(docx)
    assert "第1段落" in text
    assert "第2段落" in text
    assert "セル1-1" in text
    assert "セル2-2" in text


def test_extract_text_strips_evidence_marker(tmp_path: Path) -> None:
    docx = _make_docx(
        tmp_path / "b.docx",
        ["【甲第００１号証】", "本文の内容"],
    )
    text = extract_text_from_docx(docx)
    assert "甲第００１号証" not in text
    assert "本文の内容" in text


# ---------------------------------------------------------------------------
# truncate
# ---------------------------------------------------------------------------

def test_truncate_text_under_limit_no_change() -> None:
    text = "あいうえお" * 10  # 50 文字
    out, truncated = truncate_for_tokens(text, max_tokens=1000)
    assert out == text
    assert truncated is False


def test_truncate_text_over_limit_keeps_head_and_tail() -> None:
    head_marker = "HEAD_START"
    tail_marker = "TAIL_END"
    body = "x" * 50_000
    text = head_marker + body + tail_marker
    out, truncated = truncate_for_tokens(text, max_tokens=1000)  # max_chars = 2500
    assert truncated is True
    assert out.startswith(head_marker)
    assert out.endswith(tail_marker)
    assert "中略" in out
    assert len(out) < len(text)


# ---------------------------------------------------------------------------
# プロンプトファイル
# ---------------------------------------------------------------------------

def test_system_prompt_file_loads() -> None:
    """SPEC §10.5.5: app/prompts/extract_metadata.txt が読み込めること。"""
    prompt = load_system_prompt()
    assert "証拠説明書" in prompt
    assert "title" in prompt
    assert "confidence" in prompt


# ---------------------------------------------------------------------------
# 正常系(API モック)
# ---------------------------------------------------------------------------

def test_extract_metadata_returns_parsed_dataclass(tmp_path: Path) -> None:
    docx = _make_docx(tmp_path / "c.docx", [_long_body()])
    client = _make_client(_make_response(_VALID_JSON))
    result = extract_metadata(docx, client=client)

    assert isinstance(result, ExtractedMetadata)
    assert result.title == "業務委託契約書"
    assert result.created_date == "令和5年4月1日"
    assert result.author == "山田太郎"
    assert result.purpose == "契約締結の事実を立証"
    assert result.confidence.title == "high"
    assert result.confidence.purpose == "low"
    assert result.truncated is False


def test_extract_metadata_partial_empty_strings_allowed(tmp_path: Path) -> None:
    docx = _make_docx(tmp_path / "d.docx", [_long_body()])
    partial_json = json.dumps({
        "title": "メール記録",
        "created_date": "",
        "author": "",
        "purpose": "通信の存在を立証",
        "confidence": {
            "title": "high",
            "created_date": "low",
            "author": "low",
            "purpose": "medium",
        },
        "notes": "",
    }, ensure_ascii=False)
    client = _make_client(_make_response(partial_json))
    result = extract_metadata(docx, client=client)
    assert result.title == "メール記録"
    assert result.created_date == ""
    assert result.author == ""


def test_extract_metadata_passes_truncated_text_when_long(tmp_path: Path) -> None:
    long_body = "あ" * 100_000
    docx = _make_docx(tmp_path / "e.docx", [long_body])
    client = _make_client(_make_response(_VALID_JSON))
    result = extract_metadata(docx, max_input_tokens=1000, client=client)

    # truncate された旨が dataclass に反映される
    assert result.truncated is True
    # API に渡された user_text が原文より短いことを確認
    call_args = client.messages.create.call_args
    user_text = call_args.kwargs["messages"][0]["content"]
    assert len(user_text) < len(long_body)
    assert "中略" in user_text


# ---------------------------------------------------------------------------
# エッジケース: confidence の不正値
# ---------------------------------------------------------------------------

def test_extract_metadata_invalid_confidence_falls_back_to_low(tmp_path: Path) -> None:
    docx = _make_docx(tmp_path / "f.docx", [_long_body()])
    bad_json = json.dumps({
        "title": "契約書",
        "created_date": "令和5年",
        "author": "甲社",
        "purpose": "",
        "confidence": {
            "title": "unknown",       # 不正値
            "created_date": "very_high",  # 不正値
            "author": "high",
            "purpose": "low",
        },
        "notes": "",
    }, ensure_ascii=False)
    client = _make_client(_make_response(bad_json))
    result = extract_metadata(docx, client=client)

    assert result.confidence.title == "low"
    assert result.confidence.created_date == "low"
    assert result.confidence.author == "high"
    # フォールバックが notes に追記される
    assert "unknown" in result.notes
    assert "very_high" in result.notes


# ---------------------------------------------------------------------------
# 異常系
# ---------------------------------------------------------------------------

def test_extract_metadata_auth_error_raises_AiAuthError(tmp_path: Path) -> None:
    docx = _make_docx(tmp_path / "g.docx", [_long_body()])
    err = _raw_anthropic_error(anthropic.AuthenticationError)
    client = _make_client([err])  # 1 回で失敗。リトライ不可。
    with pytest.raises(AiAuthError):
        extract_metadata(docx, client=client)
    # リトライしないこと
    assert client.messages.create.call_count == 1


def test_extract_metadata_rate_limit_retries_then_succeeds(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = _make_docx(tmp_path / "h.docx", [_long_body()])
    monkeypatch.setattr(ai_client.time, "sleep", lambda _s: None)
    err = _raw_anthropic_error(anthropic.RateLimitError)
    client = _make_client([err, _make_response(_VALID_JSON)])
    result = extract_metadata(docx, client=client)
    assert result.title == "業務委託契約書"
    assert client.messages.create.call_count == 2


def test_extract_metadata_rate_limit_retries_then_fails(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = _make_docx(tmp_path / "i.docx", [_long_body()])
    monkeypatch.setattr(ai_client.time, "sleep", lambda _s: None)
    err = _raw_anthropic_error(anthropic.RateLimitError)
    # 初回 + 2 リトライ = 3 回すべて失敗
    client = _make_client([err, err, err])
    with pytest.raises(AiRateLimitError):
        extract_metadata(docx, client=client)
    assert client.messages.create.call_count == 3


def test_extract_metadata_network_error_retries_then_fails(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = _make_docx(tmp_path / "j.docx", [_long_body()])
    monkeypatch.setattr(ai_client.time, "sleep", lambda _s: None)
    err = _raw_anthropic_error(anthropic.APIConnectionError)
    client = _make_client([err, err, err])
    with pytest.raises(AiClientError) as exc_info:
        extract_metadata(docx, client=client)
    # ネットワークエラーは AiClientError(派生でない基底)で返る想定
    assert "ネットワーク" in str(exc_info.value)
    assert client.messages.create.call_count == 3


def test_extract_metadata_timeout_raises_AiTimeoutError(tmp_path: Path) -> None:
    docx = _make_docx(tmp_path / "k.docx", [_long_body()])
    err = _raw_anthropic_error(anthropic.APITimeoutError)
    client = _make_client([err])  # タイムアウトはリトライ対象外
    with pytest.raises(AiTimeoutError):
        extract_metadata(docx, client=client)
    assert client.messages.create.call_count == 1


def test_extract_metadata_invalid_json_raises_AiParseError(tmp_path: Path) -> None:
    docx = _make_docx(tmp_path / "l.docx", [_long_body()])
    client = _make_client(_make_response("これは JSON ではありません"))
    with pytest.raises(AiParseError):
        extract_metadata(docx, client=client)


def test_extract_metadata_missing_api_key_raises_AiAuthError(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """client を注入せず、かつ ANTHROPIC_API_KEY 未設定 → AiAuthError。"""
    docx = _make_docx(tmp_path / "m.docx", [_long_body()])
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    with pytest.raises(AiAuthError):
        extract_metadata(docx)


def test_extract_metadata_short_text_raises_value_error(tmp_path: Path) -> None:
    """SPEC §10.5.4: 100 文字未満は ValueError(API 層が 422 に変換)。"""
    docx = _make_docx(tmp_path / "n.docx", ["短い"])
    with pytest.raises(ValueError):
        extract_metadata(docx)


def test_extract_metadata_missing_file_raises_value_error(tmp_path: Path) -> None:
    with pytest.raises(ValueError):
        extract_metadata(tmp_path / "does_not_exist.docx")
