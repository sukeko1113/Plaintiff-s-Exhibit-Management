# -*- coding: utf-8 -*-
"""
分解(split)ロジックの「ページ先頭判定」に関するテスト。

section anchor として拾うのは次の 3 条件をすべて満たす段落のみ:
  (A) ページ先頭にある — 文書冒頭 / 直前段落の <w:br w:type="page"/> /
       直前段落の <w:pPr> 内 <w:sectPr> / 当該段落の <w:pPr> 内
       <w:pageBreakBefore/>
  (B) 段落冒頭がマーカーパターン (^\\s*【...】) にマッチする
  (C) 【】 で括られている

本テストは「本文中に引用された【甲第○号証】を section anchor として誤検出
しないこと」を中心に検証する。
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_BREAK

from app.split_evidence_docx import split_docx


# ---------------------------------------------------------------------------
# ヘルパ
# ---------------------------------------------------------------------------

def _save(doc: Document, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _add_paragraph_with_trailing_page_break(doc: Document, text: str) -> None:
    """段落を追加し、その末尾に <w:br w:type="page"/> を埋め込む。"""
    p = doc.add_paragraph(text)
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_page_break_before_paragraph(doc: Document, text: str) -> None:
    """段落を追加し、その <w:pPr> に <w:pageBreakBefore/> を設定する。"""
    p = doc.add_paragraph(text)
    p.paragraph_format.page_break_before = True


def _produced_names(out_dir: Path) -> list[str]:
    return sorted(p.name for p in out_dir.glob("*.docx"))


# ---------------------------------------------------------------------------
# 1. 本文中に引用された 【甲第○号証】 は section anchor として拾わない
# ---------------------------------------------------------------------------

def test_inline_citation_is_not_picked_up(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    # p1: 文書冒頭(=ページ先頭) かつ段落冒頭がマーカー
    doc.add_paragraph("【甲第１号証】")
    # p2: 通常の本文。ページ先頭ではないが、本文中に他のマーカーを引用している
    doc.add_paragraph("本文中で【甲第２号証】を引用しています。")
    _save(doc, src)

    out_dir = tmp_path / "out"
    produced = split_docx(src, out_dir, verbose=False)

    # 拾うのは p1 の 1 件のみ
    assert len(produced) == 1
    assert _produced_names(out_dir) == ["甲第００１号証.docx"]


# ---------------------------------------------------------------------------
# 2. マーカーが本文の途中にしか無い場合は ValueError
# ---------------------------------------------------------------------------

def test_only_inline_citation_raises(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    # p1: ページ先頭だがマーカー無し
    doc.add_paragraph("これは前書きの段落です。")
    # p2: ページ先頭ではなく、しかも本文中にしかマーカーが無い
    doc.add_paragraph("本文中に【甲第１号証】の引用しかありません。")
    _save(doc, src)

    out_dir = tmp_path / "out"
    with pytest.raises(ValueError):
        split_docx(src, out_dir, verbose=False)


# ---------------------------------------------------------------------------
# 3. ハード改ページ後の段落冒頭の 【甲第○号証】 は拾う
# ---------------------------------------------------------------------------

def test_hard_page_break_before_marker_is_anchor(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    # p1: 文書冒頭、マーカー
    doc.add_paragraph("【甲第１号証】")
    # p2: 本文(末尾にハード改ページ)
    _add_paragraph_with_trailing_page_break(doc, "甲第１号証の本文。")
    # p3: ページ先頭(改ページ直後)、マーカー
    doc.add_paragraph("【甲第２号証】")
    doc.add_paragraph("甲第２号証の本文。")
    _save(doc, src)

    out_dir = tmp_path / "out"
    produced = split_docx(src, out_dir, verbose=False)

    assert _produced_names(out_dir) == ["甲第００１号証.docx", "甲第００２号証.docx"]
    assert len(produced) == 2


# ---------------------------------------------------------------------------
# 4. ページ先頭の段落でも段落途中にしかマーカーが無い場合は拾わない
# ---------------------------------------------------------------------------

def test_page_top_paragraph_with_inline_marker_is_not_anchor(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    # p1: 文書冒頭、実マーカー
    doc.add_paragraph("【甲第１号証】")
    # p2: 末尾改ページ
    _add_paragraph_with_trailing_page_break(doc, "甲第１号証の本文。")
    # p3: ページ先頭だが、マーカーが段落途中にしか無い
    doc.add_paragraph("前置きの後に【甲第２号証】が書かれている。")
    _save(doc, src)

    out_dir = tmp_path / "out"
    produced = split_docx(src, out_dir, verbose=False)

    # 拾うのは p1 のみ
    assert _produced_names(out_dir) == ["甲第００１号証.docx"]
    assert len(produced) == 1


# ---------------------------------------------------------------------------
# 5. 実マーカー + 本文引用が混在しても、実マーカーだけ拾う
# ---------------------------------------------------------------------------

def test_real_markers_and_inline_citations_mixed(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    # 甲第１号証(実マーカー)
    doc.add_paragraph("【甲第１号証】")
    # 本文中で別番号を引用(セクションアンカーではない)
    _add_paragraph_with_trailing_page_break(doc, "本文中で【甲第３号証】を引用。")
    # 甲第２号証(実マーカー、ハード改ページ直後)
    doc.add_paragraph("【甲第２号証】")
    # 本文中で前の番号に言及
    doc.add_paragraph("先述の【甲第１号証】を参照。")
    _save(doc, src)

    out_dir = tmp_path / "out"
    produced = split_docx(src, out_dir, verbose=False)

    # 実マーカーは ００１ と ００２ の 2 件のみ
    assert _produced_names(out_dir) == ["甲第００１号証.docx", "甲第００２号証.docx"]
    assert len(produced) == 2


# ---------------------------------------------------------------------------
# 6. <w:pageBreakBefore/> を持つ段落もページ先頭扱い
# ---------------------------------------------------------------------------

def test_page_break_before_paragraph_is_anchor(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    # p1: 文書冒頭の前書き(マーカー無し)
    doc.add_paragraph("前書きの段落(マーカー無し)。")
    # p2: 通常の本文
    doc.add_paragraph("通常の本文。")
    # p3: <w:pageBreakBefore/> 付き → ページ先頭扱い
    _add_page_break_before_paragraph(doc, "【甲第１号証】")
    doc.add_paragraph("甲第１号証の本文。")
    _save(doc, src)

    out_dir = tmp_path / "out"
    produced = split_docx(src, out_dir, verbose=False)

    assert _produced_names(out_dir) == ["甲第００１号証.docx"]
    assert len(produced) == 1
