# -*- coding: utf-8 -*-
"""
pytest 共通フィクスチャ・ヘルパ。
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Optional

import pytest
from docx import Document

# リポジトリルートを sys.path に追加（`app` パッケージをインポートするため）
ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


def make_kogo_docx(path: Path, marker: str, body: Optional[str] = None) -> Path:
    """
    冒頭にマーカーを持つ簡易な甲号証 docx を生成する。

    body にはマーカー文字列を含めない（本文走査でマーカーが二重に拾われると
    結合後の検証が混乱するため）。
    """
    doc = Document()
    doc.add_paragraph(marker)
    doc.add_paragraph(body if body is not None else "テスト用本文。")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def collect_markers(docx_path: Path) -> list[str]:
    """
    結合後 docx から【甲第…号証(その…)】マーカーを順に拾う。
    """
    from app.kogo_normalizer import MARKER_PATTERN

    doc = Document(str(docx_path))
    out: list[str] = []
    for para in doc.paragraphs:
        for m in MARKER_PATTERN.finditer(para.text):
            # 正規化形式に整形しなおす（保存値そのまま）
            out.append(m.group(0))
    return out


def has_pagebreak(docx_path: Path, expected_count: int) -> bool:
    """
    docx 内の改ページ数を w:br[w:type='page'] で数えて検証する。
    """
    import zipfile
    from xml.etree import ElementTree as ET

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(str(docx_path)) as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    breaks = [b for b in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br")
              if b.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"]
    # 期待値以上含まれていれば OK（docxcompose が w:p 内に入れる場合あり）
    return len(breaks) >= expected_count


@pytest.fixture
def root_folder(tmp_path: Path) -> Path:
    """テスト用のルートフォルダ。空の状態を返す。"""
    return tmp_path / "case"
