#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
甲号証ファイル結合スクリプト（CLI / モジュール兼用）

複数の「甲第xxx号証.docx」ファイルを番号順に結合し、結合甲号証ファイルを生成する。
番号正規化ロジックは app.kogo_normalizer に集約してある。

使い方（CLI）::

    python -m app.merge_kogo_shoko 入力ファイル1.docx 入力ファイル2.docx ...
    python -m app.merge_kogo_shoko --indir ./個別マスタ -o 結合甲号証.docx
"""

from __future__ import annotations

import argparse
import shutil
import sys
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer

from app.kogo_normalizer import (
    KogoNumber,
    MARKER_PATTERN,
    detect_number,
)


def rewrite_marker_in_document(doc_path: Path, kogo: KogoNumber, out_path: Path) -> None:
    """文書冒頭の【甲第xxx号証】マーカーを正規化形式に書き換えて保存。"""
    doc = Document(str(doc_path))
    rewritten = False

    for para in doc.paragraphs[:10]:
        if MARKER_PATTERN.search(para.text):
            new_text = MARKER_PATTERN.sub(kogo.normalized_marker, para.text, count=1)
            if not para.runs:
                para.add_run(new_text)
            else:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            rewritten = True
            break

    if not rewritten:
        print(
            f"  [情報] {doc_path.name}: 本文マーカーは見つからず（書き換え省略）",
            file=sys.stderr,
        )

    doc.save(str(out_path))


def insert_pagebreak_at_end(doc_path: Path) -> None:
    """文書末尾に改ページを追加。"""
    doc = Document(str(doc_path))
    last_para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    run = last_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    doc.save(str(doc_path))


def merge_documents(prepared_files: List[Path], output_path: Path) -> None:
    """docxcompose で複数 docx を結合する。"""
    if not prepared_files:
        raise ValueError("結合対象ファイルがありません")

    base = Document(str(prepared_files[0]))
    composer = Composer(base)
    for f in prepared_files[1:]:
        composer.append(Document(str(f)))
    composer.save(str(output_path))


def prepare_and_merge(
    pairs: List[Tuple[KogoNumber, Path]],
    output_path: Path,
    *,
    insert_pagebreak: bool = True,
) -> None:
    """
    (KogoNumber, ファイル) のリストを受け取り、ソート→マーカー書き換え→結合。

    pairs はソート前でも構わないが、入力順をそのまま採用する場合は呼び出し側で
    並び順を整えておくこと。本関数はソートキーで安定ソートする。
    """
    if not pairs:
        raise ValueError("結合対象ファイルがありません")

    pairs_sorted = sorted(pairs, key=lambda x: x[0].sort_key)

    workdir = Path(tempfile.mkdtemp(prefix="kogo_merge_"))
    try:
        prepared: List[Path] = []
        for idx, (kogo, src) in enumerate(pairs_sorted):
            stem = kogo.normalized_filename_stem
            tmp_path = workdir / f"{idx:03d}_{stem}.docx"
            rewrite_marker_in_document(src, kogo, tmp_path)
            if insert_pagebreak and idx < len(pairs_sorted) - 1:
                insert_pagebreak_at_end(tmp_path)
            prepared.append(tmp_path)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        merge_documents(prepared, output_path)
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description="甲号証 docx ファイルを番号順に結合します。",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("inputs", nargs="*", type=Path, help="入力 docx ファイル")
    parser.add_argument("--indir", type=Path, help="入力ディレクトリ")
    parser.add_argument(
        "-o", "--output", type=Path,
        default=Path("結合甲号証.docx"),
        help="出力ファイル名",
    )
    parser.add_argument(
        "--no-pagebreak", action="store_true",
        help="各甲号証の末尾に改ページを挿入しない",
    )
    args = parser.parse_args(argv)

    inputs: List[Path] = list(args.inputs)
    if args.indir:
        inputs.extend(sorted(args.indir.glob("*.docx")))
    inputs = [p for p in inputs if not p.name.startswith("~$")]
    if not inputs:
        parser.error("入力ファイルが指定されていません")

    print("[1/3] 甲号証番号を抽出中...")
    pairs: List[Tuple[KogoNumber, Path]] = []
    for p in inputs:
        try:
            kogo = detect_number(p)
            print(f"  {p.name}  ->  {kogo.normalized_marker}")
            pairs.append((kogo, p))
        except ValueError as e:
            print(f"  [スキップ] {e}", file=sys.stderr)

    if not pairs:
        print("結合対象がありません", file=sys.stderr)
        return 1

    seen: set = set()
    for kogo, p in pairs:
        if kogo.sort_key in seen:
            print(f"  [警告] 番号重複: {kogo.normalized_marker} ({p.name})", file=sys.stderr)
        seen.add(kogo.sort_key)

    print(f"\n[2/3] 結合中 -> {args.output}")
    prepare_and_merge(pairs, args.output, insert_pagebreak=not args.no_pagebreak)
    print("[3/3] 完了。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
