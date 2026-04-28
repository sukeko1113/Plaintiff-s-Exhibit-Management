# -*- coding: utf-8 -*-
"""
甲号証リスト生成サービス。

個別マスタ または 結合甲号証 から甲号証番号を抽出し、
ルート直下の `甲号証リスト.docx` に「番号を1段落1件で縦に並べただけ」の
シンプルなリストを書き出す。

メタデータ付きの一覧テーブル（証拠説明書）とは別物である点に注意。
"""

from __future__ import annotations

import logging
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Set, Tuple

from docx import Document

from app.kogo_normalizer import (
    KogoNumber,
    MARKER_PATTERN,
    detect_number,
    extract_numbers_from_text,
)
from app.merge_service import (
    LIST_FILENAME,
    MASTER_DIRNAME,
    OUTPUT_DIRNAME,
    OUTPUT_FILENAME,
    BACKUP_SUFFIX,
    ensure_folders,
)


logger = logging.getLogger(__name__)


SOURCE_MASTER = "master"
SOURCE_COMBINED = "combined"
VALID_SOURCES = (SOURCE_MASTER, SOURCE_COMBINED)


@dataclass
class ListOutcome:
    """リスト生成結果。FastAPI の AutoListResult に詰め替える。"""

    output_path: Path
    source: str
    numbers_written: List[str] = field(default_factory=list)
    backup_created: bool = False
    warnings: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# 抽出: source="master"
# ---------------------------------------------------------------------------

def _collect_numbers_from_master(master_dir: Path) -> Tuple[List[KogoNumber], List[str]]:
    """
    個別マスタフォルダ内の docx を走査し、KogoNumber のリストを返す。

    - `~$` で始まる Word の一時ロックファイルは除外
    - 番号を抽出できないファイルは警告に記録、リストには含めない
    - 同一番号が複数ファイルに存在する場合は最初の1件のみ採用し警告に記録
    """
    warnings: List[str] = []
    numbers: List[KogoNumber] = []
    seen: Set[Tuple[int, Optional[int]]] = set()

    if not master_dir.exists():
        warnings.append(f"個別マスタフォルダが存在しません: {master_dir}")
        return numbers, warnings

    for path in sorted(master_dir.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        try:
            kogo = detect_number(path)
        except ValueError as e:
            warnings.append(str(e))
            continue
        key = (kogo.main, kogo.branch)
        if key in seen:
            warnings.append(
                f"個別マスタに重複した番号があります: {kogo.normalized_marker} ({path.name})"
            )
            continue
        seen.add(key)
        numbers.append(kogo)

    return numbers, warnings


# ---------------------------------------------------------------------------
# 抽出: source="combined"
# ---------------------------------------------------------------------------

def _collect_numbers_from_combined(combined_path: Path) -> Tuple[List[KogoNumber], List[str]]:
    """
    結合甲号証.docx を読み、本文中のマーカー（【甲第〇〇〇号証】形式）を全て拾う。

    - 結合ファイルが存在しない場合は FileNotFoundError
    - 重複は除去（最初の出現のみ採用）
    """
    warnings: List[str] = []
    numbers: List[KogoNumber] = []
    seen: Set[Tuple[int, Optional[int]]] = set()

    if not combined_path.exists():
        raise FileNotFoundError(f"結合甲号証ファイルが存在しません: {combined_path}")

    doc = Document(str(combined_path))
    for para in doc.paragraphs:
        for kogo in extract_numbers_from_text(para.text, MARKER_PATTERN):
            key = (kogo.main, kogo.branch)
            if key in seen:
                continue
            seen.add(key)
            numbers.append(kogo)

    return numbers, warnings


# ---------------------------------------------------------------------------
# 書き出し
# ---------------------------------------------------------------------------

def _write_list_docx(output_path: Path, numbers: List[KogoNumber]) -> None:
    """番号を1行1段落で縦に並べたリスト docx を出力する。"""
    doc = Document()
    for kogo in numbers:
        doc.add_paragraph(kogo.normalized_filename_stem)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# メイン
# ---------------------------------------------------------------------------

def generate_list(root_folder: Path, source: str) -> ListOutcome:
    """
    指定ソースから甲号証番号を抽出し、ルート直下に甲号証リスト.docx を書き出す。

    Parameters
    ----------
    root_folder : ルートフォルダ
    source : "master"（個別マスタから）または "combined"（結合甲号証から）

    Returns
    -------
    ListOutcome

    Raises
    ------
    ValueError
        source が "master" / "combined" のいずれでもない場合。
    FileNotFoundError
        source="combined" のときに 結合甲号証/結合甲号証.docx が存在しない場合。
    """
    if source not in VALID_SOURCES:
        raise ValueError(
            f"source は {VALID_SOURCES} のいずれかを指定してください: {source!r}"
        )

    root = Path(root_folder)
    output_path = root / LIST_FILENAME

    # ensure_folders は不在時に空の 甲号証リスト.docx を新規作成するため、
    # バックアップ対象かどうかは「呼び出し時点で既存だったか」で判定する。
    list_existed_before = output_path.exists()

    root = ensure_folders(root)

    if source == SOURCE_MASTER:
        master_dir = root / MASTER_DIRNAME
        numbers, warnings = _collect_numbers_from_master(master_dir)
    else:  # SOURCE_COMBINED
        combined_path = root / OUTPUT_DIRNAME / OUTPUT_FILENAME
        numbers, warnings = _collect_numbers_from_combined(combined_path)

    numbers.sort(key=lambda k: k.sort_key)

    # 既存リストをバックアップ（merge の出力バックアップと同じ方針）
    backup_created = False
    if list_existed_before:
        backup = output_path.with_suffix(output_path.suffix + BACKUP_SUFFIX)
        shutil.copy2(output_path, backup)
        backup_created = True
        logger.info("既存の甲号証リストをバックアップ: %s", backup)

    _write_list_docx(output_path, numbers)

    return ListOutcome(
        output_path=output_path,
        source=source,
        numbers_written=[k.normalized_filename_stem for k in numbers],
        backup_created=backup_created,
        warnings=warnings,
    )
