# -*- coding: utf-8 -*-
"""
甲号証番号の正規化・抽出ユーティリティ。

merge_kogo_shoko.py と merge_service.py の共通ロジックを集約する。
仕様書 (merge_kogo_shoko_仕様書.md) の番号正規化ルールに準拠。
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional, Tuple

from docx import Document


ZEN2HAN_DIGITS = str.maketrans("０１２３４５６７８９", "0123456789")
HAN2ZEN_DIGITS = str.maketrans("0123456789", "０１２３４５６７８９")


def to_han_digits(s: str) -> str:
    return s.translate(ZEN2HAN_DIGITS)


def to_zen_digits(s: str) -> str:
    return s.translate(HAN2ZEN_DIGITS)


def normalize_main_number(num_str: str, width: int = 3) -> str:
    """主番号を全角3桁にゼロ埋め。4桁以上はそのまま全角化。"""
    n = int(to_han_digits(num_str))
    return to_zen_digits(str(n).zfill(width))


def normalize_branch_number(num_str: str) -> str:
    """枝番を全角化（ゼロ埋めなし）。"""
    n = int(to_han_digits(num_str))
    return to_zen_digits(str(n))


# 本文マーカー: 【甲第ＮＮＮ号証(そのＭ)】
MARKER_PATTERN = re.compile(
    r"【?\s*甲\s*第?\s*"
    r"(?P<main>[0-9０-９]+)"
    r"\s*号\s*証"
    r"(?:\s*その\s*(?P<branch>[0-9０-９]+))?"
    r"\s*】?"
)

# ファイル名用（拡張子前まで）
FILENAME_PATTERN = re.compile(
    r"甲\s*第?\s*"
    r"(?P<main>[0-9０-９]+)"
    r"\s*号\s*証"
    r"(?:\s*その\s*(?P<branch>[0-9０-９]+))?"
)

# 甲号証リスト用: 段落・セルから番号だけを拾うパターン（マーカーの【】が無い書式）
LIST_PATTERN = re.compile(
    r"甲\s*第?\s*"
    r"(?P<main>[0-9０-９]+)"
    r"\s*号\s*証"
    r"(?:\s*その\s*(?P<branch>[0-9０-９]+))?"
)


@dataclass(frozen=True)
class KogoNumber:
    """甲号証番号（半角整数で保持）。"""

    main: int
    branch: Optional[int] = None

    @property
    def sort_key(self) -> Tuple[int, int]:
        return (self.main, self.branch if self.branch is not None else 0)

    @property
    def normalized_marker(self) -> str:
        main_str = normalize_main_number(str(self.main))
        if self.branch is not None:
            return f"【甲第{main_str}号証その{normalize_branch_number(str(self.branch))}】"
        return f"【甲第{main_str}号証】"

    @property
    def normalized_filename_stem(self) -> str:
        main_str = normalize_main_number(str(self.main))
        if self.branch is not None:
            return f"甲第{main_str}号証その{normalize_branch_number(str(self.branch))}"
        return f"甲第{main_str}号証"


def _build_kogo(main_raw: str, branch_raw: Optional[str]) -> KogoNumber:
    main = int(to_han_digits(main_raw))
    branch = int(to_han_digits(branch_raw)) if branch_raw else None
    return KogoNumber(main=main, branch=branch)


def extract_number_from_text(text: str, pattern: re.Pattern = MARKER_PATTERN) -> Optional[KogoNumber]:
    if not text:
        return None
    m = pattern.search(text)
    if not m:
        return None
    return _build_kogo(m.group("main"), m.group("branch"))


def extract_numbers_from_text(text: str, pattern: re.Pattern = LIST_PATTERN) -> list[KogoNumber]:
    """1つのテキストから複数の番号を全て拾う（リスト表記用）。"""
    if not text:
        return []
    return [_build_kogo(m.group("main"), m.group("branch")) for m in pattern.finditer(text)]


def extract_number_from_filename(path: Path) -> Optional[KogoNumber]:
    m = FILENAME_PATTERN.search(path.stem)
    if not m:
        return None
    return _build_kogo(m.group("main"), m.group("branch"))


def detect_number(path: Path) -> KogoNumber:
    """
    ファイル名と本文冒頭の両方から番号を検出。本文を優先。
    両方失敗した場合は ValueError を送出する。
    """
    from_doc: Optional[KogoNumber] = None
    try:
        doc = Document(str(path))
        for para in doc.paragraphs[:10]:
            from_doc = extract_number_from_text(para.text)
            if from_doc:
                break
    except Exception as e:  # pragma: no cover - 破損ファイル等
        print(f"  [警告] 本文読込失敗: {e}", file=sys.stderr)

    from_name = extract_number_from_filename(path)

    if from_doc and from_name and from_doc.sort_key != from_name.sort_key:
        print(
            f"  [警告] {path.name}: ファイル名({from_name.sort_key}) と "
            f"本文マーカー({from_doc.sort_key}) が不一致 → 本文を優先",
            file=sys.stderr,
        )

    if from_doc:
        return from_doc
    if from_name:
        return from_name
    raise ValueError(f"{path.name}: 甲号証番号を抽出できませんでした")


def iter_doc_text_blocks(doc: Document) -> Iterable[str]:
    """段落と表セルのテキストを順に列挙。"""
    for para in doc.paragraphs:
        yield para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para.text
