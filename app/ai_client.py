# -*- coding: utf-8 -*-
"""
Anthropic Claude API ラッパー(SPEC §10.5)。

個別マスタ Word ファイルからメタデータ4項目(標目・作成年月日・作成者・立証趣旨)
を AI 抽出する低レイヤモジュール。本モジュールは API 呼び出しのみを担当し、
HTTP エンドポイントとの結線は上位サービス(metadata_extractor / metadata_router)
で行う。

設計方針:
- API キーは関数呼び出し時に毎回 os.getenv("ANTHROPIC_API_KEY") を読む
  (テスト時の monkeypatch / .env 切り替えを容易にするため)
- python-dotenv の load_dotenv() は import 時に実行(冪等なので
  main.py 等で重複呼び出ししても安全)
- 例外は AiClientError 階層で表現し、上位層が HTTP ステータスにマップする
"""

from __future__ import annotations

import json
import logging
import os
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from docx import Document
from dotenv import load_dotenv

logger = logging.getLogger(__name__)

# .env を読み込む(冪等。main.py 等で重複呼び出ししても安全)
load_dotenv()


# ---------------------------------------------------------------------------
# 定数
# ---------------------------------------------------------------------------

# モデル ID(環境変数で上書き可)
DEFAULT_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6")

# 入力トークン上限(SPEC §10.5.5。環境変数で上書き可)
DEFAULT_MAX_INPUT_TOKENS = int(os.getenv("ANTHROPIC_MAX_INPUT_TOKENS", "30000"))

# タイムアウト(秒)
DEFAULT_TIMEOUT_SEC = 60

# リトライ最大回数(SPEC §10.5.5: 最大 2 回・指数バックオフ 2→4 秒)
MAX_RETRIES = 2
RETRY_BACKOFF_SEC = (2.0, 4.0)

# 極端に短い本文の閾値(SPEC §10.5.4: 100 文字未満で 422 を返す)
MIN_TEXT_LENGTH = 100

# truncate 比率: 標目・作成年月日・作成者は冒頭に、署名・日付は末尾に出やすいため、
# 冒頭 70% / 末尾 30% で切り詰める。
TRUNCATE_HEAD_RATIO = 0.7
TRUNCATE_TAIL_RATIO = 0.3

# 日本語想定のトークン換算式(概算、厳密ではない)。
# Anthropic 公式の目安に基づき 1 トークン ≒ 2.5 文字と仮定。
CHARS_PER_TOKEN_ESTIMATE = 2.5

# 出力トークン上限(JSON 4 項目 + confidence + notes に十分な余裕)
MAX_OUTPUT_TOKENS = 1024

# 本文中の証拠マーカーを除去するための正規表現(SPEC §10.5.1)。
EVIDENCE_MARKER_RE = re.compile(r"【\s*甲第[０-９0-9]+号証(?:その[０-９0-9]+)?\s*】")

# システムプロンプトファイル(SPEC §10.5.5: 外出し)
PROMPT_PATH = Path(__file__).parent / "prompts" / "extract_metadata.txt"

VALID_CONFIDENCE_VALUES = {"high", "medium", "low"}


# ---------------------------------------------------------------------------
# 例外クラス
# ---------------------------------------------------------------------------

class AiClientError(Exception):
    """AI クライアントの基底例外。"""


class AiAuthError(AiClientError):
    """認証エラー(401 / API キー未設定)。リトライ不可。"""


class AiRateLimitError(AiClientError):
    """レート制限(429)。リトライ後も失敗した場合に送出。"""


class AiTimeoutError(AiClientError):
    """タイムアウト(60 秒超)。"""


class AiParseError(AiClientError):
    """レスポンスの JSON パース失敗。"""


# ---------------------------------------------------------------------------
# データクラス
# ---------------------------------------------------------------------------

@dataclass
class ConfidenceLevels:
    title: str
    created_date: str
    author: str
    purpose: str


@dataclass
class ExtractedMetadata:
    title: str
    created_date: str
    author: str
    purpose: str
    confidence: ConfidenceLevels
    notes: str = ""
    truncated: bool = False  # 入力が truncate された場合 True(上位層でログ用)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "created_date": self.created_date,
            "author": self.author,
            "purpose": self.purpose,
            "confidence": {
                "title": self.confidence.title,
                "created_date": self.confidence.created_date,
                "author": self.confidence.author,
                "purpose": self.confidence.purpose,
            },
            "notes": self.notes,
        }


# ---------------------------------------------------------------------------
# テキスト抽出
# ---------------------------------------------------------------------------

def extract_text_from_docx(docx_path: Path) -> str:
    """python-docx で段落 + 表セルを順に平文化する。

    本文中の証拠マーカー(【甲第〇〇〇号証】)は除去する(SPEC §10.5.1)。
    """
    doc = Document(str(docx_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if text:
                    parts.append(text)

    joined = "\n".join(parts)
    # マーカー除去(SPEC §10.5.1)
    joined = EVIDENCE_MARKER_RE.sub("", joined)
    return joined.strip()


def truncate_for_tokens(text: str, max_tokens: int) -> Tuple[str, bool]:
    """入力テキストがトークン上限を超える場合、冒頭 70% + 末尾 30% で切り詰める。

    トークン換算は日本語想定の概算(1 トークン ≒ 2.5 文字)で代用する。
    戻り値: (切り詰め後テキスト, truncated フラグ)
    """
    max_chars = int(max_tokens * CHARS_PER_TOKEN_ESTIMATE)
    if len(text) <= max_chars:
        return text, False

    head_chars = int(max_chars * TRUNCATE_HEAD_RATIO)
    tail_chars = max_chars - head_chars  # 残りを末尾に
    head = text[:head_chars]
    tail = text[-tail_chars:] if tail_chars > 0 else ""
    truncated = head + "\n\n[... 中略 ...]\n\n" + tail
    return truncated, True


# ---------------------------------------------------------------------------
# プロンプト読み込み
# ---------------------------------------------------------------------------

def load_system_prompt() -> str:
    """`app/prompts/extract_metadata.txt` を読み込む。"""
    if not PROMPT_PATH.exists():
        raise AiClientError(
            f"システムプロンプトファイルが見つかりません: {PROMPT_PATH}"
        )
    return PROMPT_PATH.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# レスポンスパース
# ---------------------------------------------------------------------------

def _coerce_confidence(value: Any, notes_acc: list[str], field_name: str) -> str:
    """confidence 値を "high"/"medium"/"low" に正規化。
    不正値は "low" にフォールバックし、notes に記録する。
    """
    if isinstance(value, str) and value in VALID_CONFIDENCE_VALUES:
        return value
    notes_acc.append(
        f"confidence.{field_name} の値 {value!r} を 'low' にフォールバックしました"
    )
    return "low"


def _parse_response_json(raw: str) -> ExtractedMetadata:
    """AI 応答テキストを ExtractedMetadata に変換。

    JSON パース失敗時は AiParseError。
    confidence の不正値は "low" にフォールバック(notes に追記)。
    """
    text = raw.strip()
    # まれに markdown フェンスが付いている場合を除去
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)

    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        raise AiParseError(f"AI レスポンスの JSON パースに失敗: {e}; raw={raw!r}") from e

    if not isinstance(data, dict):
        raise AiParseError(f"AI レスポンスが JSON オブジェクトではありません: {data!r}")

    conf_raw = data.get("confidence") or {}
    if not isinstance(conf_raw, dict):
        conf_raw = {}

    notes_acc: list[str] = []
    confidence = ConfidenceLevels(
        title=_coerce_confidence(conf_raw.get("title"), notes_acc, "title"),
        created_date=_coerce_confidence(
            conf_raw.get("created_date"), notes_acc, "created_date"
        ),
        author=_coerce_confidence(conf_raw.get("author"), notes_acc, "author"),
        purpose=_coerce_confidence(conf_raw.get("purpose"), notes_acc, "purpose"),
    )

    notes = str(data.get("notes") or "")
    if notes_acc:
        suffix = " / ".join(notes_acc)
        notes = f"{notes} ({suffix})" if notes else suffix

    return ExtractedMetadata(
        title=str(data.get("title") or ""),
        created_date=str(data.get("created_date") or ""),
        author=str(data.get("author") or ""),
        purpose=str(data.get("purpose") or ""),
        confidence=confidence,
        notes=notes,
    )


# ---------------------------------------------------------------------------
# Anthropic クライアント
# ---------------------------------------------------------------------------

def _build_client(timeout_sec: float) -> Any:
    """anthropic.Anthropic クライアントを構築。

    API キーは呼び出しごとに os.getenv("ANTHROPIC_API_KEY") を参照する。
    未設定の場合は AiAuthError を即座に送出する。
    """
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise AiAuthError(
            "ANTHROPIC_API_KEY が設定されていません。.env または環境変数で設定してください。"
        )
    import anthropic  # 遅延 import(テスト時の差し替えを許容)
    return anthropic.Anthropic(api_key=api_key, timeout=timeout_sec)


def _classify_anthropic_error(exc: Exception) -> AiClientError:
    """anthropic SDK の例外を AiClientError 階層に変換する。"""
    import anthropic
    if isinstance(exc, anthropic.AuthenticationError):
        return AiAuthError(f"認証エラー: {exc}")
    if isinstance(exc, anthropic.RateLimitError):
        return AiRateLimitError(f"レート制限: {exc}")
    if isinstance(exc, anthropic.APITimeoutError):
        return AiTimeoutError(f"タイムアウト: {exc}")
    if isinstance(exc, anthropic.APIConnectionError):
        return AiClientError(f"ネットワークエラー: {exc}")
    if isinstance(exc, anthropic.APIStatusError):
        return AiClientError(f"API エラー({exc.status_code}): {exc}")
    return AiClientError(f"未分類のエラー: {exc}")


def _is_retryable(err: AiClientError) -> bool:
    """SPEC §10.5.5: リトライ対象は429・ネットワーク。認証/タイムアウト/パース失敗は不可。"""
    return isinstance(err, AiRateLimitError) or (
        type(err) is AiClientError and "ネットワーク" in str(err)
    )


def _call_with_retry(client: Any, system: str, user_text: str, model: str) -> str:
    """messages.create を呼び出し、SPEC §10.5.5 のリトライ方針(最大2回・2→4秒)を適用。

    成功時は応答テキストを返す。失敗時は適切な AiClientError 階層を送出。
    """
    last_err: Optional[AiClientError] = None
    for attempt in range(MAX_RETRIES + 1):  # 0,1,2 = 初回 + 2 リトライ
        try:
            response = client.messages.create(
                model=model,
                max_tokens=MAX_OUTPUT_TOKENS,
                system=system,
                messages=[{"role": "user", "content": user_text}],
            )
            # SDK のレスポンスから先頭テキストブロックを取得
            blocks = getattr(response, "content", None) or []
            if not blocks:
                raise AiParseError("AI レスポンスが空です")
            first = blocks[0]
            text = getattr(first, "text", None)
            if text is None:
                raise AiParseError(f"AI レスポンスにテキストが含まれません: {first!r}")
            logger.info("AI 抽出成功(attempt=%d, model=%s)", attempt, model)
            return text
        except AiParseError:
            raise
        except Exception as exc:
            err = _classify_anthropic_error(exc)
            last_err = err
            if not _is_retryable(err) or attempt >= MAX_RETRIES:
                logger.warning(
                    "AI 抽出失敗(attempt=%d, retryable=%s): %s",
                    attempt, _is_retryable(err), err,
                )
                raise err
            backoff = RETRY_BACKOFF_SEC[min(attempt, len(RETRY_BACKOFF_SEC) - 1)]
            logger.info(
                "AI 抽出リトライ(attempt=%d → %d, backoff=%.1fs): %s",
                attempt, attempt + 1, backoff, err,
            )
            time.sleep(backoff)
    # ループ脱出は理論上ありえないが、念のため
    assert last_err is not None
    raise last_err


# ---------------------------------------------------------------------------
# パブリック API
# ---------------------------------------------------------------------------

def extract_metadata(
    docx_path: Path,
    *,
    model: Optional[str] = None,
    max_input_tokens: Optional[int] = None,
    timeout_sec: float = DEFAULT_TIMEOUT_SEC,
    client: Optional[Any] = None,
) -> ExtractedMetadata:
    """個別マスタ Word ファイルから 4 項目を抽出する。

    Args:
        docx_path: 個別マスタ docx ファイルのパス
        model: モデル ID(未指定時は DEFAULT_MODEL)
        max_input_tokens: 入力トークン上限(未指定時は DEFAULT_MAX_INPUT_TOKENS)
        timeout_sec: タイムアウト秒数
        client: テスト用にモック注入できる anthropic.Anthropic 互換オブジェクト

    Raises:
        ValueError: ファイル不在、または本文が極端に短い場合(100文字未満)
        AiAuthError: API キー未設定 / 認証失敗
        AiRateLimitError: レート制限(リトライ後も失敗)
        AiTimeoutError: タイムアウト
        AiParseError: 不正な JSON 応答
        AiClientError: その他の API エラー
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise ValueError(f"対象ファイルが存在しません: {docx_path}")

    text = extract_text_from_docx(docx_path)
    if len(text) < MIN_TEXT_LENGTH:
        raise ValueError(
            f"本文が短すぎます({len(text)} 文字 < {MIN_TEXT_LENGTH}): {docx_path}"
        )

    used_model = model or DEFAULT_MODEL
    used_max_tokens = max_input_tokens if max_input_tokens is not None else DEFAULT_MAX_INPUT_TOKENS
    truncated_text, truncated = truncate_for_tokens(text, used_max_tokens)
    if truncated:
        logger.info(
            "入力テキストを truncate(原文=%d文字 → %d文字, max_tokens=%d)",
            len(text), len(truncated_text), used_max_tokens,
        )

    system_prompt = load_system_prompt()
    used_client = client if client is not None else _build_client(timeout_sec)

    raw = _call_with_retry(used_client, system_prompt, truncated_text, used_model)
    result = _parse_response_json(raw)
    result.truncated = truncated
    return result
