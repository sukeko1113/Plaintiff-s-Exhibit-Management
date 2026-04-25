"""仕様 §9 table_builder テスト。"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.core.table_builder import (
    build_summary_doc,
    parse_metadata,
    preview_metadata,
)


def _make_with_meta(path: Path, marker: str, meta: dict[str, str], body: str = '') -> None:
    doc = Document()
    doc.add_paragraph(marker)
    for k in ['標目', '作成年月日', '作成者', '立証趣旨']:
        if k in meta:
            doc.add_paragraph(f'{k}: {meta[k]}')
    if body:
        doc.add_paragraph('')
        doc.add_paragraph(body)
    doc.save(str(path))


def _make_without_meta(path: Path, marker: str, body: str) -> None:
    doc = Document()
    doc.add_paragraph(marker)
    doc.add_paragraph(body)
    doc.save(str(path))


def test_parse_metadata_complete(tmp_path):
    f = tmp_path / '甲第００１号証.docx'
    _make_with_meta(f, '【甲第１号証】', {
        '標目': '報告書',
        '作成年月日': '令和8年1月1日',
        '作成者': '〇〇株式会社',
        '立証趣旨': '事実Ａを証明する。',
    }, body='報告書本文…')

    meta = parse_metadata(f)
    assert meta == {
        '標目': '報告書',
        '作成年月日': '令和8年1月1日',
        '作成者': '〇〇株式会社',
        '立証趣旨': '事実Ａを証明する。',
    }


def test_parse_metadata_partial(tmp_path):
    f = tmp_path / '甲第００２号証.docx'
    _make_with_meta(f, '【甲第２号証】', {
        '標目': '契約書',
    }, body='契約書本文…')

    meta = parse_metadata(f)
    assert meta == {'標目': '契約書'}


def test_parse_metadata_none(tmp_path):
    f = tmp_path / '甲第００３号証.docx'
    _make_without_meta(f, '【甲第３号証】', '本文だけ。')
    assert parse_metadata(f) == {}


def test_parse_metadata_full_width_colon(tmp_path):
    """全角コロン「：」も区切りとして許容。"""
    f = tmp_path / '甲第００４号証.docx'
    doc = Document()
    doc.add_paragraph('【甲第４号証】')
    doc.add_paragraph('標目：契約書')
    doc.add_paragraph('作成者：被告')
    doc.save(str(f))

    assert parse_metadata(f) == {'標目': '契約書', '作成者': '被告'}


def test_build_summary_doc_with_metadata(tmp_path):
    f1 = tmp_path / '甲第００１号証.docx'
    f2 = tmp_path / '甲第００２号証.docx'
    _make_with_meta(f1, '【甲第１号証】', {
        '標目': '報告書', '作成年月日': '令和8年1月1日',
        '作成者': '原告', '立証趣旨': '事実Ａを証明する。',
    })
    _make_without_meta(f2, '【甲第２号証】', '本文')

    out = tmp_path / 'summary.docx'
    doc = build_summary_doc([f1, f2])
    doc.save(str(out))

    re_opened = Document(str(out))
    assert any('証拠説明書' in p.text for p in re_opened.paragraphs)
    table = re_opened.tables[0]
    headers = [c.text for c in table.rows[0].cells]
    assert headers == ['号証', '標目', '作成年月日', '作成者', '立証趣旨']

    row1 = [c.text for c in table.rows[1].cells]
    assert row1 == ['甲第1号証', '報告書', '令和8年1月1日', '原告', '事実Ａを証明する。']

    row2 = [c.text for c in table.rows[2].cells]
    # メタなし → 標目はファイル名から、他は空
    assert row2 == ['甲第2号証', '甲第2号証', '', '', '']


def test_metadata_map_overrides_file(tmp_path):
    """metadata_map で渡されたものが本文メタより優先。"""
    f = tmp_path / '甲第００１号証.docx'
    _make_with_meta(f, '【甲第１号証】', {'標目': 'ファイル本文の標目'})

    doc = build_summary_doc([f], metadata_map={
        '甲第００１号証': {'標目': 'API指定標目', '立証趣旨': 'API指定趣旨'},
    })
    out = tmp_path / 'summary.docx'
    doc.save(str(out))

    table = Document(str(out)).tables[0]
    row = [c.text for c in table.rows[1].cells]
    assert row[1] == 'API指定標目'
    assert row[4] == 'API指定趣旨'


def test_preview_metadata(tmp_path):
    f = tmp_path / '甲第００１号証.docx'
    _make_with_meta(f, '【甲第１号証】', {'標目': 'タイトル'})
    rows = preview_metadata([f])
    assert len(rows) == 1
    assert rows[0]['label'] == '甲第００１号証'
    assert rows[0]['display_label'] == '甲第1号証'
    assert rows[0]['標目'] == 'タイトル'
