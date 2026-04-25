"""テスト用結合 docx の生成（仕様 §14.2）。

各号証を `docxcompose` で結合した状態の `combined_*.docx` を生成する。
各号証の冒頭段落に【甲第N号証】マーカーを置き、splitter のテスト入力にする。
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docxcompose.composer import Composer

FIXTURES_DIR = Path(__file__).parent / 'fixtures'


def _make_section_doc(marker: str, body_lines: list[str]) -> Document:
    """1 号証分の Document を作成。先頭段落にマーカーを置く。"""
    doc = Document()
    doc.add_paragraph(marker)
    for line in body_lines:
        doc.add_paragraph(line)
    return doc


def _make_section_doc_with_table(marker: str) -> Document:
    doc = Document()
    doc.add_paragraph(marker)
    doc.add_paragraph('表を含む証拠書類:')
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = '項目'
    headers[1].text = '値'
    headers[2].text = '備考'
    data = table.rows[1].cells
    data[0].text = '日付'
    data[1].text = '令和8年1月1日'
    data[2].text = '記載例'
    return doc


def _make_section_doc_with_image(marker: str) -> Document:
    doc = Document()
    doc.add_paragraph(marker)
    doc.add_paragraph('画像を含む証拠書類:')
    doc.add_picture(io.BytesIO(_make_minimal_png()), width=Inches(0.5))
    return doc


def _make_minimal_png() -> bytes:
    """有効な 1x1 グレースケール PNG をバイト列として生成。"""
    import struct
    import zlib

    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack('>I', len(data)) + tag + data + struct.pack('>I', zlib.crc32(tag + data))

    ihdr = struct.pack('>IIBBBBB', 1, 1, 8, 0, 0, 0, 0)
    idat = zlib.compress(b'\x00\x00')
    return b'\x89PNG\r\n\x1a\n' + chunk(b'IHDR', ihdr) + chunk(b'IDAT', idat) + chunk(b'IEND', b'')


def _combine_all(docs: list[Document], output: Path) -> None:
    if not docs:
        raise ValueError('docs is empty')
    output.parent.mkdir(parents=True, exist_ok=True)
    composer = Composer(docs[0])
    for d in docs[1:]:
        composer.append(d)
    composer.save(str(output))


def generate_combined_simple() -> Path:
    """5 件の号証（甲第1〜5号証）を結合。"""
    docs = [
        _make_section_doc(f'【甲第{i}号証】', [f'本文サンプル {i}-1', f'本文サンプル {i}-2'])
        for i in range(1, 6)
    ]
    out = FIXTURES_DIR / 'combined_simple.docx'
    _combine_all(docs, out)
    return out


def generate_combined_with_branch() -> Path:
    """5 件のうち甲第3号証は枝番（その1 / その2）を持つ。"""
    docs = [
        _make_section_doc('【甲第1号証】', ['本文 1']),
        _make_section_doc('【甲第2号証】', ['本文 2']),
        _make_section_doc('【甲第3号証】', ['本文 3 主']),
        _make_section_doc('【甲第3号証その1】', ['本文 3-枝1']),
        _make_section_doc('【甲第3号証その2】', ['本文 3-枝2']),
        _make_section_doc('【甲第4号証】', ['本文 4']),
    ]
    out = FIXTURES_DIR / 'combined_with_branch.docx'
    _combine_all(docs, out)
    return out


def generate_combined_with_table() -> Path:
    docs = [
        _make_section_doc_with_table(f'【甲第{i}号証】')
        for i in range(1, 4)
    ]
    out = FIXTURES_DIR / 'combined_with_table.docx'
    _combine_all(docs, out)
    return out


def generate_combined_with_image() -> Path:
    docs = [
        _make_section_doc_with_image(f'【甲第{i}号証】')
        for i in range(1, 4)
    ]
    out = FIXTURES_DIR / 'combined_with_image.docx'
    _combine_all(docs, out)
    return out


def generate_case_sample() -> Path:
    """本文・表・脚注に甲第〇〇号証を散りばめた案件ファイル。"""
    doc = Document()
    doc.add_heading('訴状（サンプル）', level=1)
    doc.add_paragraph('原告は被告に対し、甲第1号証に基づき以下の事実を主張する。')
    doc.add_paragraph('さらに、甲第２号証および【甲第3号証】を参照されたい。')
    doc.add_paragraph('予備的に甲第１号証の2を提出する。')

    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '号証'
    table.rows[0].cells[1].text = '内容'
    table.rows[1].cells[0].text = '甲第5号証'
    table.rows[1].cells[1].text = '報告書（参照: 甲第10号証その1）'

    doc.add_paragraph('普通の文章で、甲号証は出てきません。')

    out = FIXTURES_DIR / 'case_sample.docx'
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def generate_all() -> dict[str, Path]:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    return {
        'combined_simple': generate_combined_simple(),
        'combined_with_branch': generate_combined_with_branch(),
        'combined_with_table': generate_combined_with_table(),
        'combined_with_image': generate_combined_with_image(),
        'case_sample': generate_case_sample(),
    }


if __name__ == '__main__':
    paths = generate_all()
    for name, p in paths.items():
        print(f'{name}: {p} ({p.stat().st_size} bytes)')
