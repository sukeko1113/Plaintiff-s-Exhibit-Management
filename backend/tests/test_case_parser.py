"""仕様 §10 case_parser テスト。"""

from __future__ import annotations

import pytest
from docx import Document

from backend.core.case_parser import extract_koshou_from_case
from backend.tests import generate_fixtures


@pytest.fixture(scope='module', autouse=True)
def _ensure_fixtures():
    generate_fixtures.generate_all()


def test_extract_from_sample_case():
    """generate_fixtures の case_sample.docx に含まれる全マーカーを抽出。

    fixture には以下が含まれる:
    - 本文: 甲第1号証, 甲第２号証, 【甲第3号証】, 甲第１号証の2
    - 表中: 甲第5号証, 甲第10号証その1
    - ノイズ: 普通の文章
    """
    src = generate_fixtures.FIXTURES_DIR / 'case_sample.docx'
    labels = extract_koshou_from_case(src)
    assert labels == [
        '甲第００１号証',
        '甲第００１号証その２',
        '甲第００２号証',
        '甲第００３号証',
        '甲第００５号証',
        '甲第０１０号証その１',
    ]


def test_extract_dedupes(tmp_path):
    case = tmp_path / 'case.docx'
    doc = Document()
    doc.add_paragraph('甲第1号証')
    doc.add_paragraph('甲第１号証')  # 同一の表記ゆれ
    doc.add_paragraph('甲第001号証')  # 同一
    doc.save(str(case))

    assert extract_koshou_from_case(case) == ['甲第００１号証']


def test_extract_missing_file(tmp_path):
    with pytest.raises(FileNotFoundError):
        extract_koshou_from_case(tmp_path / 'nonexistent.docx')


def test_extract_empty_doc(tmp_path):
    case = tmp_path / 'empty.docx'
    Document().save(str(case))
    assert extract_koshou_from_case(case) == []


def test_extract_otsugou_is_ignored(tmp_path):
    """乙号証は対象外。"""
    case = tmp_path / 'mix.docx'
    doc = Document()
    doc.add_paragraph('甲第1号証は引用するが、乙第1号証は対象外。')
    doc.save(str(case))

    assert extract_koshou_from_case(case) == ['甲第００１号証']
