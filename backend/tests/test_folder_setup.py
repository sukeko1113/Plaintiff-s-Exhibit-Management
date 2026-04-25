"""仕様 §11 のテスト。"""

from __future__ import annotations

import pytest

from backend.core.folder_setup import REQUIRED_FILES, REQUIRED_FOLDERS, setup_root


def test_setup_creates_all(tmp_path):
    result = setup_root(tmp_path)

    assert result['root'] == str(tmp_path)
    for folder in REQUIRED_FOLDERS:
        assert (tmp_path / folder).is_dir()
        assert folder in result['created']
    for fname in REQUIRED_FILES:
        assert (tmp_path / fname).is_file()
        assert fname in result['created']
    assert result['existed'] == []


def test_setup_idempotent(tmp_path):
    setup_root(tmp_path)
    result2 = setup_root(tmp_path)

    assert result2['created'] == []
    for name in REQUIRED_FOLDERS + REQUIRED_FILES:
        assert name in result2['existed']


def test_setup_rejects_missing(tmp_path):
    missing = tmp_path / 'does_not_exist'
    with pytest.raises(FileNotFoundError):
        setup_root(missing)


def test_setup_rejects_file(tmp_path):
    f = tmp_path / 'file.txt'
    f.write_text('x')
    with pytest.raises(NotADirectoryError):
        setup_root(f)


def test_setup_does_not_overwrite_existing_list(tmp_path):
    """既存の甲号証リスト.docx は触らない。"""
    from docx import Document

    list_path = tmp_path / '甲号証リスト.docx'
    doc = Document()
    doc.add_paragraph('既存のラベル')
    doc.save(str(list_path))

    setup_root(tmp_path)

    reopened = Document(str(list_path))
    assert any('既存のラベル' in p.text for p in reopened.paragraphs)
