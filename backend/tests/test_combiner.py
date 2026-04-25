"""仕様 §7 / §14.1 combiner テスト。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from backend.core.combiner import combine_to_evidence_pack
from backend.core.splitter import split_combined_to_master
from backend.tests import generate_fixtures


@pytest.fixture(scope='module', autouse=True)
def _ensure_fixtures():
    generate_fixtures.generate_all()


def _split_then_combine(tmp_path: Path, src_name: str) -> tuple[list[Path], Path]:
    """fixtures の結合 docx を分割 → 個別マスタを再結合した出力パスを返す。"""
    src = generate_fixtures.FIXTURES_DIR / src_name
    master_dir = tmp_path / 'master'
    individual_files = split_combined_to_master(src, master_dir)

    out_path = tmp_path / 'combined.docx'
    result_path = combine_to_evidence_pack(individual_files, out_path)
    return individual_files, result_path


def test_combine_simple(tmp_path):
    files, combined = _split_then_combine(tmp_path, 'combined_simple.docx')
    assert len(files) == 5
    assert combined.exists()

    # 結合結果が docx として開けて、5 件分のマーカーをすべて含む
    doc = Document(str(combined))
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    for i in range(1, 6):
        assert f'【甲第{i}号証】' in full_text


def test_combine_with_branch_sorted(tmp_path):
    files, combined = _split_then_combine(tmp_path, 'combined_with_branch.docx')
    assert len(files) == 6

    doc = Document(str(combined))
    paragraphs = [p.text for p in doc.paragraphs]
    full = '\n'.join(paragraphs)

    # マーカーすべて含む
    for marker in ['【甲第1号証】', '【甲第2号証】', '【甲第3号証】',
                   '【甲第3号証その1】', '【甲第3号証その2】', '【甲第4号証】']:
        assert marker in full

    # 順序確認: 3号証 → その1 → その2 → 4号証
    idx_main = full.index('【甲第3号証】')
    idx_b1 = full.index('【甲第3号証その1】')
    idx_b2 = full.index('【甲第3号証その2】')
    idx_4 = full.index('【甲第4号証】')
    assert idx_main < idx_b1 < idx_b2 < idx_4


def test_combine_unsorted_input_is_sorted(tmp_path):
    """入力リストが順不同でも、出力は番号順になる。"""
    src = generate_fixtures.FIXTURES_DIR / 'combined_simple.docx'
    master_dir = tmp_path / 'master'
    individual_files = split_combined_to_master(src, master_dir)

    # 逆順で渡す
    reversed_files = list(reversed(individual_files))
    out_path = tmp_path / 'combined.docx'
    combine_to_evidence_pack(reversed_files, out_path)

    doc = Document(str(out_path))
    full = '\n'.join(p.text for p in doc.paragraphs)
    indices = [full.index(f'【甲第{i}号証】') for i in range(1, 6)]
    assert indices == sorted(indices), '番号順になっていない'


def test_combine_empty_raises(tmp_path):
    out_path = tmp_path / 'combined.docx'
    with pytest.raises(ValueError):
        combine_to_evidence_pack([], out_path)


def test_combine_creates_parent_dir(tmp_path):
    files, _ = _split_then_combine(tmp_path, 'combined_simple.docx')
    out_path = tmp_path / 'nested' / 'subdir' / 'out.docx'
    combine_to_evidence_pack(files, out_path)
    assert out_path.exists()
