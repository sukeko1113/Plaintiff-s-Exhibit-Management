"""splitter / combiner / case_parser / list_builder の統合テスト（v02）。"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from backend.core.backup import backup_paths, list_generations
from backend.core.case_parser import extract_labels_from_case_file
from backend.core.combiner import auto_filename, combine_files, ensure_docx_extension
from backend.core.folder_setup import setup_root_folder, get_list_path
from backend.core.list_builder import (
    auto_create_from_master,
    labels_from_combined_file,
    parse_list_file,
    write_list_file,
)
from backend.core.splitter import preview_split, split_combined_file
from backend.tests.generate_fixtures import (
    make_case_fixture,
    make_combined_fixture_bare,
    make_combined_fixture_bracketed,
    make_list_fixture,
    make_master_files,
)


def test_split_bracketed_marker(tmp_path: Path) -> None:
    combined = tmp_path / 'combined.docx'
    make_combined_fixture_bracketed(combined)
    out_dir = tmp_path / 'master'

    extracted = split_combined_file(combined, out_dir)
    assert [e.label for e in extracted] == [
        '甲第００１号証', '甲第００２号証', '甲第００２号証その１',
    ]
    for e in extracted:
        assert (out_dir / e.filename).exists()


def test_split_bare_marker_fallback(tmp_path: Path) -> None:
    combined = tmp_path / 'combined_bare.docx'
    make_combined_fixture_bare(combined)
    out_dir = tmp_path / 'master'

    extracted = split_combined_file(combined, out_dir)
    assert [e.label for e in extracted] == [
        '甲第００１号証', '甲第００２号証', '甲第００２号証その１',
    ]


def test_preview_split_dry_run(tmp_path: Path) -> None:
    combined = tmp_path / 'combined.docx'
    make_combined_fixture_bracketed(combined)
    assert preview_split(combined) == [
        '甲第００１号証', '甲第００２号証', '甲第００２号証その１',
    ]


def test_labels_from_combined_file(tmp_path: Path) -> None:
    combined = tmp_path / 'combined.docx'
    make_combined_fixture_bracketed(combined)
    assert labels_from_combined_file(combined) == [
        '甲第００１号証', '甲第００２号証', '甲第００２号証その１',
    ]


def test_combine_then_split_round_trip(tmp_path: Path) -> None:
    master = tmp_path / 'master'
    make_master_files(master)
    combined = tmp_path / 'out' / 'combined.docx'

    labels = ['甲第００１号証', '甲第００２号証', '甲第００２号証その１']
    result = combine_files(master, labels, combined)
    assert result.missing == []
    assert result.used == labels
    assert combined.exists()

    extracted = split_combined_file(combined, tmp_path / 'extracted')
    assert [e.label for e in extracted] == labels

    body_text = '\n'.join(p.text for p in Document(str(combined)).paragraphs)
    assert '甲第１号証の中身' in body_text
    assert '甲第２号証の枝番' in body_text


def test_combine_reports_missing_labels(tmp_path: Path) -> None:
    master = tmp_path / 'master'
    make_master_files(master)
    combined = tmp_path / 'out' / 'combined.docx'

    labels = ['甲第００１号証', '甲第９９９号証']
    result = combine_files(master, labels, combined)
    assert result.missing == ['甲第９９９号証']
    assert result.used == ['甲第００１号証']


def test_auto_filename_includes_root_and_kind() -> None:
    name = auto_filename('令和8年（ワ）第131号', '結合')
    assert name.startswith('令和8年（ワ）第131号_甲号証_結合_')
    assert name.endswith('.docx')


def test_ensure_docx_extension() -> None:
    assert ensure_docx_extension('foo') == 'foo.docx'
    assert ensure_docx_extension('foo.docx') == 'foo.docx'


def test_case_parser_extracts_from_body_and_table(tmp_path: Path) -> None:
    case = tmp_path / 'case.docx'
    make_case_fixture(case)
    labels = extract_labels_from_case_file(str(case))
    assert labels == [
        '甲第００１号証', '甲第００２号証その１', '甲第００３号証', '甲第０１２号証',
    ]


def test_split_raises_when_no_split_points(tmp_path: Path) -> None:
    empty = tmp_path / 'empty.docx'
    Document().save(str(empty))
    with pytest.raises(ValueError):
        split_combined_file(empty, tmp_path / 'master')


def test_setup_creates_backup_dir_and_summary(tmp_path: Path) -> None:
    result = setup_root_folder(str(tmp_path))
    assert result.ok
    for name in ('甲号証リスト.docx', '個別マスタ', '結合甲号証', '_backup'):
        assert (tmp_path / name).exists()
    assert result.summary == {
        'master_count': 0,
        'list_label_count': 1,  # コメント行 1 行分
        'combined_files_count': 0,
        'backup_generations': 0,
    }


def test_parse_list_returns_ignored_lines(tmp_path: Path) -> None:
    setup_root_folder(str(tmp_path))
    list_path = get_list_path(str(tmp_path))
    make_list_fixture(list_path)
    parsed = parse_list_file(str(tmp_path))
    assert parsed.labels == ['甲第００１号証', '甲第００２号証', '甲第００２号証その１']
    assert len(parsed.ignored_lines) == 1
    assert 'このファイル' in parsed.ignored_lines[0].text


def test_write_list_creates_backup_when_overwriting(tmp_path: Path) -> None:
    setup_root_folder(str(tmp_path))
    write_list_file(str(tmp_path), ['甲第００１号証'])
    write_list_file(str(tmp_path), ['甲第００２号証'])
    gens = list_generations(str(tmp_path))
    assert len(gens) >= 1


def test_backup_keeps_only_max_generations(tmp_path: Path) -> None:
    setup_root_folder(str(tmp_path))
    list_path = get_list_path(str(tmp_path))
    for _ in range(15):
        backup_paths(str(tmp_path), [list_path])
    gens = list_generations(str(tmp_path))
    assert len(gens) <= 10


def test_auto_create_from_master_lists_master(tmp_path: Path) -> None:
    setup_root_folder(str(tmp_path))
    make_master_files(tmp_path / '個別マスタ')
    labels, _ = auto_create_from_master(str(tmp_path))
    assert labels == ['甲第００１号証', '甲第００２号証', '甲第００２号証その１']
