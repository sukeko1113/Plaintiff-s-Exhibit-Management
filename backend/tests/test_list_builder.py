"""仕様 §8 list_builder テスト。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from backend.core.list_builder import (
    build_from_combined,
    build_from_master,
    read_list,
    write_list,
)
from backend.tests import generate_fixtures


@pytest.fixture(scope='module', autouse=True)
def _ensure_fixtures():
    generate_fixtures.generate_all()


def _create_master_files(master_dir: Path, names: list[str]) -> None:
    master_dir.mkdir(parents=True, exist_ok=True)
    for n in names:
        Document().save(str(master_dir / n))


def test_build_from_master_reads_filenames(tmp_path):
    master = tmp_path / 'master'
    _create_master_files(master, [
        '甲第００１号証.docx',
        '甲第００３号証その１.docx',
        '甲第００２号証.docx',
        '甲第００３号証.docx',
    ])
    labels = build_from_master(master)
    assert labels == [
        '甲第００１号証',
        '甲第００２号証',
        '甲第００３号証',
        '甲第００３号証その１',
    ]


def test_build_from_master_skips_non_docx(tmp_path):
    master = tmp_path / 'master'
    master.mkdir()
    Document().save(str(master / '甲第００１号証.docx'))
    (master / '甲第００２号証.txt').write_text('x')
    (master / '~$甲第００３号証.docx').write_text('lock')
    (master / 'README').write_text('x')

    labels = build_from_master(master)
    assert labels == ['甲第００１号証']


def test_build_from_master_normalizes_filenames(tmp_path):
    """ファイル名の表記ゆれ（半角桁数バラバラ）も正規化される。"""
    master = tmp_path / 'master'
    _create_master_files(master, [
        '甲第1号証.docx',
        '甲第02号証.docx',
        '甲第１２号証その3.docx',
    ])
    labels = build_from_master(master)
    assert labels == [
        '甲第００１号証',
        '甲第００２号証',
        '甲第０１２号証その３',
    ]


def test_build_from_combined(tmp_path):
    src = generate_fixtures.FIXTURES_DIR / 'combined_with_branch.docx'
    labels = build_from_combined([src])
    assert labels == [
        '甲第００１号証',
        '甲第００２号証',
        '甲第００３号証',
        '甲第００３号証その１',
        '甲第００３号証その２',
        '甲第００４号証',
    ]


def test_build_from_combined_dedupes_across_files(tmp_path):
    src1 = generate_fixtures.FIXTURES_DIR / 'combined_simple.docx'
    src2 = generate_fixtures.FIXTURES_DIR / 'combined_with_branch.docx'
    labels = build_from_combined([src1, src2])
    # simple は 1〜5、branch は 1〜4 + 3その1/2 → 重複排除すれば下記
    assert labels == [
        '甲第００１号証',
        '甲第００２号証',
        '甲第００３号証',
        '甲第００３号証その１',
        '甲第００３号証その２',
        '甲第００４号証',
        '甲第００５号証',
    ]


def test_write_and_read_roundtrip(tmp_path):
    list_path = tmp_path / '甲号証リスト.docx'
    labels = ['甲第００２号証', '甲第００１号証', '甲第００１号証']  # 重複・順不同
    write_list(list_path, labels)

    assert list_path.exists()
    assert read_list(list_path) == ['甲第００１号証', '甲第００２号証']


def test_read_list_handles_format_variants(tmp_path):
    """リストに表記ゆれや空段落・コメントが混じっていても正規化する。"""
    list_path = tmp_path / '甲号証リスト.docx'
    doc = Document()
    doc.add_paragraph('甲第1号証')          # 半角・桁不揃い
    doc.add_paragraph('')                   # 空段落
    doc.add_paragraph('普通のコメント')     # ラベル以外
    doc.add_paragraph('甲第１２号証その3') # 全角混在
    doc.save(str(list_path))

    assert read_list(list_path) == ['甲第００１号証', '甲第０１２号証その３']
