"""個別マスタ → 結合甲号証 への結合（仕様書 v02 §7.6）。"""
from __future__ import annotations

import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from .normalizer import label_to_filename


@dataclass
class CombineResult:
    output_path: Path
    missing: List[str]
    used: List[str]
    backup_generation: Optional[Path] = None


def auto_filename(root_name: str, kind: str) -> str:
    """``<root_name>_甲号証_<kind>_YYYYMMDD-HHMMSS.docx`` 形式のファイル名を作る。"""
    stamp = datetime.now().strftime('%Y%m%d-%H%M%S')
    safe_root = root_name.strip() or 'output'
    return f'{safe_root}_甲号証_{kind}_{stamp}.docx'


def ensure_docx_extension(name: str) -> str:
    return name if name.lower().endswith('.docx') else f'{name}.docx'


def _has_leading_page_break(para) -> bool:
    for run in para.runs:
        for br in run._element.findall(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
    return False


def _ensure_starts_with_page_break(path: Path) -> None:
    doc = Document(str(path))
    if not doc.paragraphs:
        para = doc.add_paragraph()
    else:
        para = doc.paragraphs[0]
    if _has_leading_page_break(para):
        return
    target_run = para.runs[0] if para.runs else para.add_run()
    target_run.add_break(WD_BREAK.PAGE)
    doc.save(str(path))


def combine_files(
    master_dir: Path,
    labels: List[str],
    output_path: Path,
    head_doc: Optional[Path] = None,
) -> CombineResult:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    available: List[tuple[str, Path]] = []
    missing: List[str] = []
    for label in labels:
        candidate = master_dir / label_to_filename(label)
        if candidate.exists():
            available.append((label, candidate))
        else:
            missing.append(label)

    if not available and head_doc is None:
        raise ValueError('結合できる個別マスタファイルが 1 件もありません。')

    work_dir = output_path.parent / f'.{output_path.stem}_tmp'
    if work_dir.exists():
        shutil.rmtree(work_dir)
    work_dir.mkdir(parents=True, exist_ok=False)

    used: List[str] = []
    try:
        if head_doc is not None:
            base_path = work_dir / 'head.docx'
            shutil.copyfile(head_doc, base_path)
            files_to_append = available
        else:
            first_label, first_src = available[0]
            base_path = work_dir / f'000_{first_label}.docx'
            shutil.copyfile(first_src, base_path)
            files_to_append = available[1:]
            used.append(first_label)

        base_doc = Document(str(base_path))
        composer = Composer(base_doc)

        for idx, (label, src) in enumerate(files_to_append, start=1):
            tmp = work_dir / f'{idx:03d}_{label}.docx'
            shutil.copyfile(src, tmp)
            _ensure_starts_with_page_break(tmp)
            composer.append(Document(str(tmp)))
            used.append(label)

        composer.save(str(output_path))
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    return CombineResult(output_path=output_path, missing=missing, used=used)
