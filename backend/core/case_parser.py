"""案件ファイル（申立書 .docx）から使用号証を抽出する（仕様書 §7.7）。"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable, List

from docx import Document

from .normalizer import KOSHOU_PATTERN, koshou_sort_key, parse_match


def _iter_text_blocks(doc: Document) -> Iterable[str]:
    for para in doc.paragraphs:
        if para.text:
            yield para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        yield para.text


def extract_labels_from_case_file(case_file_path: str) -> List[str]:
    """案件ファイルから甲号証ラベルを抽出して、重複排除＋ソートして返す。"""
    path = Path(case_file_path)
    if not path.exists():
        raise FileNotFoundError(f'案件ファイルが見つかりません: {path}')
    if path.suffix.lower() != '.docx':
        raise ValueError(f'案件ファイルは .docx 形式である必要があります: {path}')

    doc = Document(str(path))
    seen: set[str] = set()
    labels: List[str] = []
    for text in _iter_text_blocks(doc):
        for match in KOSHOU_PATTERN.finditer(text):
            normalized = parse_match(match)
            if normalized is None:
                continue
            if normalized in seen:
                continue
            seen.add(normalized)
            labels.append(normalized)
    labels.sort(key=koshou_sort_key)
    return labels
