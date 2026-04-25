"""ルートフォルダの初期化（仕様書 v02 §4.1, §7.1）。"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List

from docx import Document

from .backup import BACKUP_DIR, list_generations

LIST_FILENAME = '甲号証リスト.docx'
MASTER_DIR = '個別マスタ'
COMBINED_DIR = '結合甲号証'
LIST_TEMPLATE_COMMENT = (
    '（このファイルは甲号証管理アプリが管理します。1 行 1 号証ラベルで記入してください。'
    '例: 甲第００１号証）'
)


@dataclass
class SetupResult:
    ok: bool
    messages: List[str]
    summary: Dict[str, int] = field(default_factory=dict)


def _ensure_dir(path: Path, label: str, messages: List[str]) -> None:
    if path.exists():
        if not path.is_dir():
            raise ValueError(f'「{label}」がフォルダではありません: {path}')
        messages.append(f'✅ 「{label}」フォルダを確認しました。')
    else:
        path.mkdir(parents=True, exist_ok=False)
        messages.append(f'✅ 「{label}」フォルダを作成しました。')


def _ensure_list_file(path: Path, messages: List[str]) -> None:
    if path.exists():
        if not path.is_file():
            raise ValueError(f'「{LIST_FILENAME}」がファイルではありません: {path}')
        messages.append(f'✅ 「{LIST_FILENAME}」を確認しました。')
        return
    doc = Document()
    doc.add_paragraph(LIST_TEMPLATE_COMMENT)
    doc.save(str(path))
    messages.append(f'✅ 「{LIST_FILENAME}」を作成しました。')


def setup_root_folder(root_path: str) -> SetupResult:
    """ルートフォルダ配下に必須のファイル／フォルダを作成し、サマリーを返す。"""
    root = Path(root_path).expanduser()
    if not root.exists():
        raise FileNotFoundError(f'指定のルートフォルダが存在しません: {root}')
    if not root.is_dir():
        raise NotADirectoryError(f'指定のパスがフォルダではありません: {root}')

    messages: List[str] = []
    _ensure_list_file(root / LIST_FILENAME, messages)
    _ensure_dir(root / MASTER_DIR, MASTER_DIR, messages)
    _ensure_dir(root / COMBINED_DIR, COMBINED_DIR, messages)
    _ensure_dir(root / BACKUP_DIR, BACKUP_DIR, messages)

    summary = collect_summary(root)
    return SetupResult(ok=True, messages=messages, summary=summary)


def collect_summary(root_path: str | Path) -> Dict[str, int]:
    """サマリ情報（個別マスタ件数・リスト行数・結合甲号証ファイル数・バックアップ世代数）。"""
    root = Path(root_path)
    master = root / MASTER_DIR
    combined = root / COMBINED_DIR
    list_path = root / LIST_FILENAME

    master_count = 0
    if master.exists():
        master_count = sum(
            1
            for p in master.iterdir()
            if p.is_file() and p.suffix.lower() == '.docx' and not p.name.endswith('.bak.docx')
        )

    combined_count = 0
    if combined.exists():
        combined_count = sum(
            1
            for p in combined.iterdir()
            if p.is_file() and p.suffix.lower() == '.docx' and not p.name.endswith('.bak.docx')
        )

    list_count = 0
    if list_path.exists():
        try:
            doc = Document(str(list_path))
            list_count = sum(1 for para in doc.paragraphs if para.text.strip())
        except Exception:
            list_count = 0

    backup_generations = len(list_generations(root))

    return {
        'master_count': master_count,
        'list_label_count': list_count,
        'combined_files_count': combined_count,
        'backup_generations': backup_generations,
    }


def get_master_dir(root_path: str) -> Path:
    return Path(root_path) / MASTER_DIR


def get_combined_dir(root_path: str) -> Path:
    return Path(root_path) / COMBINED_DIR


def get_list_path(root_path: str) -> Path:
    return Path(root_path) / LIST_FILENAME


def get_backup_dir(root_path: str) -> Path:
    return Path(root_path) / BACKUP_DIR


def list_combined_files(root_path: str) -> List[str]:
    """結合甲号証フォルダ内の .docx ファイル名（更新日時の新しい順）。"""
    combined = get_combined_dir(root_path)
    if not combined.exists():
        return []
    files = [
        p
        for p in combined.iterdir()
        if p.is_file() and p.suffix.lower() == '.docx' and not p.name.endswith('.bak.docx')
    ]
    files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return [p.name for p in files]


def list_combined_files_detailed(root_path: str) -> List[Dict[str, object]]:
    from datetime import datetime as _dt
    combined = get_combined_dir(root_path)
    if not combined.exists():
        return []
    out: List[Dict[str, object]] = []
    for p in combined.iterdir():
        if not p.is_file() or p.suffix.lower() != '.docx':
            continue
        if p.name.endswith('.bak.docx'):
            continue
        stat = p.stat()
        out.append(
            {
                'filename': p.name,
                'size': stat.st_size,
                'modified_at': _dt.fromtimestamp(stat.st_mtime).isoformat(timespec='seconds'),
            }
        )
    out.sort(key=lambda d: d['modified_at'], reverse=True)
    return out
