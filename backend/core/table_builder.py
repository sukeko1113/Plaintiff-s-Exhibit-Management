"""証拠説明書テーブルの自動生成（仕様書 §7.8）。

UI で渡された table_data を優先利用し、未指定の行は対応する個別マスタの
本文から「標目／作成年月日／作成者」を完全自動抽出（推測）して埋める。
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .normalizer import label_to_filename


@dataclass
class EvidenceRow:
    label: str
    title: str = ''
    date: str = ''
    author: str = ''
    purpose: str = ''


_DATE_PATTERNS = [
    re.compile(r'(令和|平成|昭和)\s*([0-9０-９元]+)\s*年\s*([0-9０-９]+)\s*月\s*([0-9０-９]+)\s*日'),
    re.compile(r'([12][0-9０-９]{3})\s*年\s*([0-9０-９]+)\s*月\s*([0-9０-９]+)\s*日'),
]

_AUTHOR_PATTERNS = [
    re.compile(r'(?:作成者|差出人|発信者|発行者|発行|作成)[:：]\s*(.+)'),
    re.compile(r'^(.+?)\s+(?:殿|様|御中)$'),
]

_TITLE_BLOCKLIST = re.compile(r'^甲第[０-９]{3}号証(?:その[０-９]+)?$')


def _first_meaningful_paragraphs(doc: Document, limit: int = 30) -> List[str]:
    out: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        out.append(text)
        if len(out) >= limit:
            break
    return out


def _guess_title(paragraphs: List[str]) -> str:
    for text in paragraphs:
        if _TITLE_BLOCKLIST.match(text):
            continue
        if len(text) <= 60:
            return text
    return ''


def _guess_date(paragraphs: List[str]) -> str:
    for text in paragraphs:
        for pat in _DATE_PATTERNS:
            m = pat.search(text)
            if m:
                return m.group(0)
    return ''


def _guess_author(paragraphs: List[str]) -> str:
    for text in paragraphs:
        for pat in _AUTHOR_PATTERNS:
            m = pat.search(text)
            if m:
                return m.group(1).strip()
    return ''


def auto_extract_row(label: str, master_dir: Path) -> EvidenceRow:
    """個別マスタファイルの先頭付近のテキストから候補値を抽出する。"""
    path = master_dir / label_to_filename(label)
    if not path.exists():
        return EvidenceRow(label=label)
    try:
        doc = Document(str(path))
    except Exception:
        return EvidenceRow(label=label)
    paragraphs = _first_meaningful_paragraphs(doc)
    return EvidenceRow(
        label=label,
        title=_guess_title(paragraphs),
        date=_guess_date(paragraphs),
        author=_guess_author(paragraphs),
        purpose='',
    )


def merge_rows(
    labels: List[str],
    master_dir: Path,
    overrides: Optional[List[Dict[str, str]]] = None,
) -> List[EvidenceRow]:
    """labels の順で行を作成。overrides が来た要素はそちらを優先する。"""
    overrides_by_label: Dict[str, Dict[str, str]] = {}
    if overrides:
        for entry in overrides:
            lbl = entry.get('label')
            if lbl:
                overrides_by_label[lbl] = entry

    rows: List[EvidenceRow] = []
    for label in labels:
        if label in overrides_by_label:
            o = overrides_by_label[label]
            rows.append(
                EvidenceRow(
                    label=label,
                    title=o.get('title', ''),
                    date=o.get('date', ''),
                    author=o.get('author', ''),
                    purpose=o.get('purpose', ''),
                )
            )
        else:
            rows.append(auto_extract_row(label, master_dir))
    return rows


def build_evidence_table_doc(rows: List[EvidenceRow], output_path: Path) -> Path:
    """証拠説明書テーブルだけを含む docx を作成する。"""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('証 拠 説 明 書')
    run.bold = True
    run.font.size = Pt(16)

    headers = ['号証', '標目', '作成年月日', '作成者', '立証趣旨']
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, name in enumerate(headers):
        hdr_cells[i].text = name
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        cells[0].text = row.label
        cells[1].text = row.title
        cells[2].text = row.date
        cells[3].text = row.author
        cells[4].text = row.purpose

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
