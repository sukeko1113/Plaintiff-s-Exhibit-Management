"""甲号証リスト.docx の自動作成・解析（仕様書 v02 §4.4, §7.4, §7.5）。"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from docx import Document

from .backup import backup_paths
from .folder_setup import (
    LIST_TEMPLATE_COMMENT,
    get_combined_dir,
    get_list_path,
    get_master_dir,
)
from .normalizer import (
    filename_to_label,
    koshou_sort_key,
    normalize_koshou_strict,
)
from .splitter import find_split_points


@dataclass
class IgnoredLine:
    line: int
    text: str


@dataclass
class ParsedList:
    labels: List[str]
    ignored_lines: List[IgnoredLine]


def labels_from_master(root_path: str) -> List[str]:
    master = get_master_dir(root_path)
    if not master.exists():
        return []
    labels: List[str] = []
    for path in master.iterdir():
        if not path.is_file() or path.suffix.lower() != '.docx':
            continue
        if path.name.endswith('.bak.docx'):
            continue
        label = filename_to_label(path.name)
        if label:
            labels.append(label)
    seen: set[str] = set()
    deduped: List[str] = []
    for label in sorted(labels, key=koshou_sort_key):
        if label in seen:
            continue
        seen.add(label)
        deduped.append(label)
    return deduped


def labels_from_combined_file(combined_path: Path) -> List[str]:
    doc = Document(str(combined_path))
    return [p.label for p in find_split_points(doc)]


def write_list_file(
    root_path: str, labels: List[str], with_comment: bool = True
) -> tuple[Path, Optional[Path]]:
    """甲号証リスト.docx を上書き保存し、(保存先, バックアップ世代) を返す。"""
    list_path = get_list_path(root_path)
    backup_generation = backup_paths(root_path, [list_path]) if list_path.exists() else None

    doc = Document()
    if with_comment:
        doc.add_paragraph(LIST_TEMPLATE_COMMENT)
    for label in labels:
        doc.add_paragraph(label)
    doc.save(str(list_path))
    return list_path, backup_generation


def parse_list_file(root_path: str) -> ParsedList:
    """甲号証リスト.docx を読み、ラベルと無視行を返す。"""
    list_path = get_list_path(root_path)
    if not list_path.exists():
        return ParsedList(labels=[], ignored_lines=[])
    doc = Document(str(list_path))
    labels: List[str] = []
    ignored: List[IgnoredLine] = []
    seen: set[str] = set()
    for index, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue
        normalized = normalize_koshou_strict(text)
        if normalized is None:
            ignored.append(IgnoredLine(line=index, text=text))
            continue
        if normalized in seen:
            continue
        seen.add(normalized)
        labels.append(normalized)
    return ParsedList(labels=labels, ignored_lines=ignored)


def auto_create_from_master(root_path: str) -> tuple[List[str], Optional[Path]]:
    labels = labels_from_master(root_path)
    _, backup = write_list_file(root_path, labels)
    return labels, backup


def auto_create_from_combined(
    root_path: str, combined_filename: str
) -> tuple[List[str], Optional[Path]]:
    combined_path = get_combined_dir(root_path) / combined_filename
    if not combined_path.exists():
        raise FileNotFoundError(f'結合甲号証ファイルが見つかりません: {combined_path}')
    labels = labels_from_combined_file(combined_path)
    _, backup = write_list_file(root_path, labels)
    return labels, backup
