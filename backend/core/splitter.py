"""結合甲号証 → 個別マスタ への分解（仕様書 v02 §6, §7.2）。

v02 では【甲第xxx号証】マーカーを最優先で検出する。1 件も見つからなければ、
括弧なしの単独行マーカー（§6.5 MARKER_BARE_STRICT_PATTERN）にフォールバックする。

安全性のため、元ファイルを丸ごとコピーしてから「自分の担当範囲外の段落」を
削除する方式を採用する（スタイル・ヘッダ・フッタ・画像参照を保持）。
"""
from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn

from .normalizer import (
    is_bare_marker,
    is_bracketed_marker,
    label_to_filename,
    normalize_koshou_strict,
)


@dataclass
class SplitPoint:
    paragraph_index: int
    label: str


@dataclass
class ExtractedFile:
    label: str
    filename: str


def _has_page_break_or_section(para) -> bool:
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        if pPr.find(qn('w:pageBreakBefore')) is not None:
            return True
        if pPr.find(qn('w:sectPr')) is not None:
            return True
    for run in para.runs:
        for br in run._element.findall(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
    return False


def _previous_is_separator(paragraphs, index: int) -> bool:
    """この段落の直前に段落区切り（空段落・改ページ・セクション切替）があるかを返す。

    自身が改ページ持ちの段落である場合も区切りと見なす（結合時に挿入された改ページ）。
    """
    if index <= 0:
        return True
    if _has_page_break_or_section(paragraphs[index]):
        return True
    prev = paragraphs[index - 1]
    if not prev.text.strip():
        return True
    return _has_page_break_or_section(prev)


def find_split_points(doc: Document) -> List[SplitPoint]:
    """v02 §6 のロジックで甲号証の開始位置を返す。"""
    paragraphs = doc.paragraphs

    bracket_points: List[SplitPoint] = []
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if is_bracketed_marker(text):
            normalized = normalize_koshou_strict(text)
            if normalized:
                bracket_points.append(SplitPoint(paragraph_index=i, label=normalized))
    if bracket_points:
        return bracket_points

    fallback: List[SplitPoint] = []
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text or not is_bare_marker(text):
            continue
        if not _previous_is_separator(paragraphs, i):
            continue
        normalized = normalize_koshou_strict(text)
        if normalized:
            fallback.append(SplitPoint(paragraph_index=i, label=normalized))
    return fallback


def _delete_paragraph_element(para) -> None:
    p = para._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _slice_document(src: Path, dst: Path, start_index: int, end_index: Optional[int]) -> None:
    shutil.copyfile(src, dst)
    doc = Document(str(dst))
    paragraphs = list(doc.paragraphs)
    total = len(paragraphs)
    end = total if end_index is None else end_index

    for idx in range(total - 1, end - 1, -1):
        if idx >= total:
            continue
        _delete_paragraph_element(paragraphs[idx])

    for idx in range(start_index - 1, -1, -1):
        _delete_paragraph_element(paragraphs[idx])

    doc.save(str(dst))


def split_combined_file(combined_path: Path, output_dir: Path) -> List[ExtractedFile]:
    """結合甲号証ファイルを分解して個別ファイルとして保存する。"""
    if not combined_path.exists():
        raise FileNotFoundError(f'結合甲号証ファイルが見つかりません: {combined_path}')
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document(str(combined_path))
    points = find_split_points(doc)
    if not points:
        raise ValueError(
            'マーカー（【甲第xxx号証】 または単独行の号証ラベル）が見つかりません。'
            ' 結合甲号証ファイルの先頭マーカー表記を確認してください。'
        )

    extracted: List[ExtractedFile] = []
    for i, point in enumerate(points):
        end_index = points[i + 1].paragraph_index if i + 1 < len(points) else None
        filename = label_to_filename(point.label)
        dst = output_dir / filename
        _slice_document(combined_path, dst, point.paragraph_index, end_index)
        extracted.append(ExtractedFile(label=point.label, filename=filename))

    return extracted


def preview_split(combined_path: Path) -> List[str]:
    """dry-run 用に、分解された場合のラベル一覧だけを返す。"""
    if not combined_path.exists():
        raise FileNotFoundError(f'結合甲号証ファイルが見つかりません: {combined_path}')
    doc = Document(str(combined_path))
    points = find_split_points(doc)
    return [p.label for p in points]
